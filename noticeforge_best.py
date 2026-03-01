# -*- coding: utf-8 -*-
"""
NoticeForge Core Logic v6.0 (Ultimate: 法令・通知・マニュアル3層対応)
  v6.0: 文書タイプ自動判別（法令/通知/マニュアル）・法令条文構造認識・相互参照マップ
  v5.4: OCR品質スコア・構造化概要・改廃追跡・法令抽出・時系列ソート・差分レポート
"""
from __future__ import annotations
import os, sys, re, json, time, hashlib, csv, subprocess, shutil, html as _html
from dataclasses import dataclass, asdict
from typing import Dict, List, Tuple, Optional, Callable

# キャッシュバージョン: 概要生成ロジックを変更した場合はインクリメントする
# → 古いキャッシュの概要が新ロジックと不整合になるのを防止
_CACHE_VERSION = 5

# Tesseract バイナリの候補パス（複数のインストール場所に対応）
_TESSERACT_CANDIDATES = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    r"C:\Users\Public\Tesseract-OCR\tesseract.exe",
]

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

TESSERACT_AVAILABLE = False
try:
    import pytesseract
    from PIL import Image
    # バイナリを自動検出（インストール場所が異なる環境に対応）
    _found_tesseract: Optional[str] = None
    for _tc in _TESSERACT_CANDIDATES:
        if os.path.isfile(_tc):
            _found_tesseract = _tc
            break
    if _found_tesseract is None:
        # PATH上にある場合（Linux / Mac / PATH追加済みのWindows）
        import shutil as _shutil
        if _shutil.which("tesseract"):
            _found_tesseract = "tesseract"
    if _found_tesseract:
        pytesseract.pytesseract.tesseract_cmd = _found_tesseract
        TESSERACT_AVAILABLE = True
except Exception:
    TESSERACT_AVAILABLE = False

try:
    from docx import Document
except Exception:
    Document = None

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.utils import get_column_letter
except Exception:
    openpyxl = None

try:
    import xlrd
except Exception:
    xlrd = None

def _setup_xdw_dll_path():
    """XDWAPI.dllのディレクトリをPythonのDLL検索パスに追加する。

    DocuWorks / DocuWorks Viewer Light がインストールされていても、
    XDWAPI.dllがPATHに含まれていないとxdwlibがインポートできない。
    レジストリとglobでインストールパスを自動検索し、
    os.add_dll_directory()（Python 3.8+）とPATHの両方に追加する。
    """
    if not sys.platform.startswith("win"):
        return
    try:
        import winreg
        import glob as _glob

        dll_dirs: List[str] = []

        reg_keys = [
            (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Fuji Xerox\DocuWorks"),
            (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\Fuji Xerox\DocuWorks"),
            (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\FUJIFILM\DocuWorks"),
            (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\FUJIFILM\DocuWorks"),
            # DocuWorks Viewer Light 専用キー（Viewer Lightはフルとは別キーに登録されることがある）
            (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\FUJIFILM\DocuWorks Viewer Light"),
            (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\FUJIFILM\DocuWorks Viewer Light"),
            (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Fuji Xerox\DocuWorks Viewer Light"),
            (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\Fuji Xerox\DocuWorks Viewer Light"),
            (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Fujitsu\DocuWorks"),
            (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\Fujitsu\DocuWorks"),
            (winreg.HKEY_CURRENT_USER,  r"SOFTWARE\Fuji Xerox\DocuWorks"),
            (winreg.HKEY_CURRENT_USER,  r"SOFTWARE\FUJIFILM\DocuWorks"),
            (winreg.HKEY_CURRENT_USER,  r"SOFTWARE\FUJIFILM\DocuWorks Viewer Light"),
        ]
        for hive, key_path in reg_keys:
            try:
                key = winreg.OpenKey(hive, key_path)
                for vname in ("InstallPath", "Path", "Install_Dir", ""):
                    try:
                        val, _ = winreg.QueryValueEx(key, vname)
                        d = str(val).strip()
                        for rel in ("", "bin", "Program", "DocuWorks", r"DocuWorks Viewer Light", r"DocuWorks Viewer Light 10"):
                            base = os.path.join(d, rel) if rel else d
                            if os.path.isfile(os.path.join(base, "XDWAPI.dll")) and base not in dll_dirs:
                                dll_dirs.append(base)
                    except Exception:
                        continue
            except Exception:
                continue

        for pattern in [
            # DocuWorks フル版
            r"C:\Program Files\Fuji Xerox\DocuWorks\XDWAPI.dll",
            r"C:\Program Files\FUJIFILM\DocuWorks\XDWAPI.dll",
            r"C:\Program Files (x86)\Fuji Xerox\DocuWorks\XDWAPI.dll",
            r"C:\Program Files (x86)\FUJIFILM\DocuWorks\XDWAPI.dll",
            # DocuWorks Viewer Light 専用パス（バージョン10等）
            r"C:\Program Files\Fuji Xerox\DocuWorks Viewer Light\XDWAPI.dll",
            r"C:\Program Files\FUJIFILM\DocuWorks Viewer Light\XDWAPI.dll",
            r"C:\Program Files\FUJIFILM\DocuWorks Viewer Light 10\XDWAPI.dll",
            r"C:\Program Files (x86)\Fuji Xerox\DocuWorks Viewer Light\XDWAPI.dll",
            r"C:\Program Files (x86)\FUJIFILM\DocuWorks Viewer Light\XDWAPI.dll",
            r"C:\Program Files (x86)\FUJIFILM\DocuWorks Viewer Light 10\XDWAPI.dll",
            # ワイルドカードでバージョン番号付きフォルダも拾う
            r"C:\Program Files\*\DocuWorks\XDWAPI.dll",
            r"C:\Program Files (x86)\*\DocuWorks\XDWAPI.dll",
            r"C:\Program Files\FUJIFILM\*\XDWAPI.dll",
            r"C:\Program Files\Fuji Xerox\*\XDWAPI.dll",
            r"C:\Program Files (x86)\FUJIFILM\*\XDWAPI.dll",
            r"C:\Program Files (x86)\Fuji Xerox\*\XDWAPI.dll",
            r"C:\Program Files\*\*\XDWAPI.dll",
            r"C:\Program Files (x86)\*\*\XDWAPI.dll",
            # システム全域
            r"C:\Windows\System32\XDWAPI.dll",
            r"C:\Windows\SysWOW64\XDWAPI.dll",
        ]:
            if "*" in pattern:
                for found in _glob.glob(pattern):
                    d = os.path.dirname(found)
                    if d not in dll_dirs:
                        dll_dirs.append(d)
            elif os.path.isfile(pattern):
                d = os.path.dirname(pattern)
                if d not in dll_dirs:
                    dll_dirs.append(d)

        for d in dll_dirs:
            if hasattr(os, "add_dll_directory"):
                try:
                    os.add_dll_directory(d)
                except Exception:
                    pass
            cur_path = os.environ.get("PATH", "")
            if d.lower() not in cur_path.lower():
                os.environ["PATH"] = d + os.pathsep + cur_path
    except Exception:
        pass

_setup_xdw_dll_path()

try:
    import xdwlib
    XDWLIB_AVAILABLE = True
except Exception:
    XDWLIB_AVAILABLE = False
_WIN_NO_CONSOLE: dict = (
    {"creationflags": 0x08000000} if sys.platform.startswith("win") else {}
)

def _build_xdw2text_candidates() -> List[str]:
    """xdw2text.exeの候補パスを構築する。
    レジストリ自動検出 → Program Files全体スキャン → 固定パスの順で探す。
    TokiwaWorks / DocuWorks Viewer / 任意のバージョンを自動検出できる。"""
    candidates: List[str] = ["xdw2text"]  # まずPATH上を探す

    if sys.platform.startswith("win"):
        # ── 方法①: Windowsレジストリを検索してインストールパスを自動検出 ──
        try:
            import winreg
            reg_keys = [
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Fuji Xerox\DocuWorks"),
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\Fuji Xerox\DocuWorks"),
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\FUJIFILM\DocuWorks"),
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\FUJIFILM\DocuWorks"),
                # DocuWorks Viewer Light 専用キー
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\FUJIFILM\DocuWorks Viewer Light"),
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\FUJIFILM\DocuWorks Viewer Light"),
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Fuji Xerox\DocuWorks Viewer Light"),
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\Fuji Xerox\DocuWorks Viewer Light"),
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Fujitsu\DocuWorks"),
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\Fujitsu\DocuWorks"),
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\TokiwaWorks\TokiwaWorks"),
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\TokiwaWorks\TokiwaWorks"),
                (winreg.HKEY_CURRENT_USER,  r"SOFTWARE\Fuji Xerox\DocuWorks"),
                (winreg.HKEY_CURRENT_USER,  r"SOFTWARE\FUJIFILM\DocuWorks"),
                (winreg.HKEY_CURRENT_USER,  r"SOFTWARE\TokiwaWorks\TokiwaWorks"),
            ]
            for hive, key_path in reg_keys:
                try:
                    key = winreg.OpenKey(hive, key_path)
                    for value_name in ("InstallPath", "Path", "Install_Dir", ""):
                        try:
                            install_path, _ = winreg.QueryValueEx(key, value_name)
                            base = str(install_path)
                            rel_candidates = [
                                "xdw2text.exe",
                                os.path.join("bin", "xdw2text.exe"),
                                os.path.join("Program", "xdw2text.exe"),
                                os.path.join("DocuWorks", "xdw2text.exe"),
                                os.path.join("DocuWorks Viewer Light", "xdw2text.exe"),
                                os.path.join("DocuWorks Viewer Light 10", "xdw2text.exe"),
                            ]
                            for rel in rel_candidates:
                                exe = os.path.join(base, rel)
                                if os.path.isfile(exe) and exe not in candidates:
                                    candidates.insert(1, exe)
                        except Exception:
                            continue
                except Exception:
                    continue
        except Exception:
            pass

        # ── 方法②: C:\Program Files 以下を glob で自動スキャン ──
        # TokiwaWorks / DocuWorks Viewer など任意のインストール先を検出できる
        try:
            import glob as _glob
            for pattern in [
                r"C:\Program Files\*\xdw2text.exe",
                r"C:\Program Files (x86)\*\xdw2text.exe",
                r"C:\Program Files\*\*\xdw2text.exe",
                r"C:\Program Files (x86)\*\*\xdw2text.exe",
                r"C:\Program Files\*\*\*\xdw2text.exe",
                r"C:\Program Files (x86)\*\*\*\xdw2text.exe",
            ]:
                for found in _glob.glob(pattern):
                    if found not in candidates:
                        candidates.insert(1, found)
        except Exception:
            pass

        # ── 方法③: 固定パス（フォールバック） ──
        candidates += [
            r"C:\Program Files\Fuji Xerox\DocuWorks\xdw2text.exe",
            r"C:\Program Files (x86)\Fuji Xerox\DocuWorks\xdw2text.exe",
            r"C:\Program Files\FUJIFILM\DocuWorks\xdw2text.exe",
            r"C:\Program Files (x86)\FUJIFILM\DocuWorks\xdw2text.exe",
            # DocuWorks Viewer Light 10 専用パス
            r"C:\Program Files\Fuji Xerox\DocuWorks Viewer Light\xdw2text.exe",
            r"C:\Program Files (x86)\Fuji Xerox\DocuWorks Viewer Light\xdw2text.exe",
            r"C:\Program Files\FUJIFILM\DocuWorks Viewer Light\xdw2text.exe",
            r"C:\Program Files\FUJIFILM\DocuWorks Viewer Light 10\xdw2text.exe",
            r"C:\Program Files (x86)\FUJIFILM\DocuWorks Viewer Light\xdw2text.exe",
            r"C:\Program Files (x86)\FUJIFILM\DocuWorks Viewer Light 10\xdw2text.exe",
            r"C:\Program Files\TokiwaWorks\xdw2text.exe",
            r"C:\Program Files (x86)\TokiwaWorks\xdw2text.exe",
            r"C:\Program Files\DocuWorks\xdw2text.exe",
            r"C:\Program Files (x86)\DocuWorks\xdw2text.exe",
        ]
    return candidates

# 起動時に候補リストを構築（レジストリも参照）
XDW2TEXT_CANDIDATES = _build_xdw2text_candidates()
# 一度見つかった実行ファイルのパスをキャッシュ（ファイルごとに7回試行しなくて済む）
_XDW2TEXT_PATH: Optional[str] = None

def _build_xdoc2txt_candidates() -> List[str]:
    """xdoc2txt.exeの候補パスを構築する。
    xdoc2txtはDocuWorks(.xdw)を含む多形式に対応した無料テキスト抽出ツール。
    https://ebstudio.info/home/xdoc2txt.html"""
    candidates: List[str] = ["xdoc2txt"]  # まずPATH上を探す
    if sys.platform.startswith("win"):
        try:
            import glob as _glob
            for pattern in [
                r"C:\Program Files\xdoc2txt\xdoc2txt.exe",
                r"C:\Program Files (x86)\xdoc2txt\xdoc2txt.exe",
                r"C:\Program Files\*\xdoc2txt.exe",
                r"C:\Program Files (x86)\*\xdoc2txt.exe",
                r"C:\tools\xdoc2txt\xdoc2txt.exe",
                r"C:\xdoc2txt\xdoc2txt.exe",
            ]:
                for found in _glob.glob(pattern):
                    if found not in candidates:
                        candidates.insert(1, found)
        except Exception:
            pass
        candidates += [
            r"C:\Program Files\xdoc2txt\xdoc2txt.exe",
            r"C:\Program Files (x86)\xdoc2txt\xdoc2txt.exe",
        ]
    return candidates

XDOC2TXT_CANDIDATES = _build_xdoc2txt_candidates()
_XDOC2TXT_PATH: Optional[str] = None

DEFAULTS: Dict[str, object] = {
    "min_chars_mainbody": 400, # 基準を少し甘くして抽出漏れを防止
    "max_depth": 30,
    "summary_chars": 900,
    "main_attach_split_keywords": [r"^\s*別添", r"^\s*別紙", r"^\s*【別添】", r"^\s*【別紙】", r"^\s*【参考】", r"^\s*記\s*$"],
    "bind_bytes_limit": 15 * 1024 * 1024,
    "use_ocr": False,
}

FACILITY_TAGS: Dict[str, List[str]] = {
    "製造所":         [r"製造所"],
    "屋外タンク貯蔵所": [r"屋外タンク貯蔵所", r"浮屋根", r"固定屋根", r"アニュラ", r"タンク底板?", r"泡放射", r"防油堤"],
    "屋内貯蔵所":     [r"屋内貯蔵所"],
    "地下タンク貯蔵所": [r"地下タンク貯蔵所", r"FRPタンク", r"漏えい検知"],
    "簡易タンク貯蔵所": [r"簡易タンク貯蔵所"],
    "移動タンク貯蔵所": [r"移動タンク貯蔵所", r"タンクローリー"],
    "給油取扱所":     [r"給油取扱所", r"計量機", r"ノズル", r"\bSS\b", r"サービスステーション", r"ガソリンスタンド"],
    "販売取扱所":     [r"販売取扱所"],
    "移送取扱所":     [r"移送取扱所", r"荷卸し", r"荷積み"],
    "一般取扱所":     [r"一般取扱所", r"塗装", r"洗浄", r"混合", r"充填", r"乾燥"],
    # ※「共通」タグは廃止。施設が特定できない場合はタグなしとする。
}

# ── 業務タグ（危険物行政専門家・消防職員との協議に基づく6分類） ──
# 廃止: 「運用解釈・Q&A」（曖昧）「事故・漏えい・火災」「消火・防災」（重複・細分化しすぎ）
# 追加: 「法令・改正」
# 統合: 事故・漏えい・火災 + 消火・防災 → 「事故・応急対応」
WORK_TAGS: Dict[str, List[str]] = {
    "法令・改正":   [r"改正", r"省令", r"告示", r"政令改正", r"規則改正",
                    r"公布", r"施行", r"通達", r"法改正", r"条の改正"],
    "申請・届出":   [r"許可", r"届出", r"申請", r"変更", r"仮使用",
                    r"完成検査", r"予防規程", r"承認", r"届書", r"様式"],
    "技術基準・設備": [r"技術基準", r"構造", r"設備", r"配管",
                    r"保有空地", r"耐震", r"腐食", r"漏えい検知", r"防油堤", r"通気管"],
    "立入検査・点検": [r"立入", r"検査", r"指導", r"是正", r"改善", r"点検", r"報告"],
    "事故・応急対応": [r"事故", r"漏えい", r"流出", r"火災", r"爆発",
                    r"消火", r"警報", r"緊急遮断", r"避難", r"防災",
                    r"消火設備", r"再発防止", r"災害"],
    "保安体制・教育": [r"保安監督", r"危険物保安監督者", r"保安統括",
                    r"教育", r"訓練", r"体制", r"責任者"],
}

@dataclass
class Record:
    relpath: str
    ext: str
    size: int
    mtime: float
    sha1: str
    method: str
    pages: Optional[int]
    text_chars: int
    needs_review: bool
    reason: str
    title_guess: str
    date_guess: str
    issuer_guess: str
    summary: str
    tags_facility: List[str]
    tags_work: List[str]
    tag_evidence: Dict[str, List[str]]
    out_txt: str
    full_text_for_bind: str = ""
    doc_type: str = "通知"             # 文書タイプ（"法令" / "通知" / "マニュアル"）
    ocr_quality: float = 1.0          # OCR品質スコア（0.0〜1.0）
    related_laws: List[str] = None     # 関連法令（「政令第○条」等）
    amendments: List[str] = None       # 改廃情報（「〜を一部改正」等）
    date_sort_key: str = ""            # 日付のソートキー（YYYYMMDD形式）

    def __post_init__(self):
        if self.related_laws is None:
            self.related_laws = []
        if self.amendments is None:
            self.amendments = []

def get_safe_path(path: str) -> str:
    """Windowsの260文字制限(MAX_PATH)を突破するための安全なパス変換"""
    abs_path = os.path.abspath(path)
    if sys.platform.startswith("win") and not abs_path.startswith("\\\\?\\"):
        return "\\\\?\\" + abs_path
    return abs_path

def extract_pdf(path: str, use_ocr: bool) -> Tuple[str, Optional[int], str]:
    if not fitz: return "", None, "pymupdf_missing"
    text_parts = []
    method = "pdf_text"
    try:
        doc = fitz.open(get_safe_path(path))
        pages = doc.page_count
        for i in range(pages):
            page = doc.load_page(i)
            page_text = page.get_text("text") or ""
            # PyMuPDF が日本語文字間にスペースを挿入する問題を修正
            # （行をまたぐ改行は残し、同一行内の不要スペースのみ除去）
            # 日本語文字間の不要スペースを除去（数字↔日本語間は箇条書き番号等で意味があるため除去しない）
            page_text = re.sub(r'([ぁ-んァ-ン一-龥])[ \t]+([ぁ-んァ-ン一-龥])', r'\1\2', page_text)
            # OCR判断:
            #   use_ocr=True → 50文字未満のページにOCR（手動指定モード）
            #   use_ocr=False → 10文字未満の極端に空なページにのみ自動OCR（画像PDF自動検出）
            ocr_trigger = 50 if use_ocr else 10
            if len(page_text.strip()) < ocr_trigger and TESSERACT_AVAILABLE:
                try:
                    pix = page.get_pixmap(dpi=200)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_text = pytesseract.image_to_string(img, lang="jpn")
                    # OCRテキストの日本語文字間スペースを除去
                    ocr_text = re.sub(r'([ぁ-んァ-ン一-龥])\s+([ぁ-んァ-ン一-龥])', r'\1\2', ocr_text)
                    if ocr_text.strip():
                        # 完全に空だったページはOCR結果で置換、テキストがあった場合は追記
                        page_text = ocr_text if len(page_text.strip()) < 10 else page_text + "\n" + ocr_text
                        method = "pdf_ocr" if use_ocr else "pdf_ocr_auto"
                except Exception:
                    pass
            text_parts.append(page_text)
        doc.close()
        return "\n".join(text_parts), pages, method
    except Exception as e:
        return "", None, f"pdf_err:{e.__class__.__name__}"

def extract_docx(path: str) -> Tuple[str, str]:
    if not Document: return "", "docx_missing"
    try:
        doc = Document(get_safe_path(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    parts.append("| " + " | ".join(cells) + " |")
        return "\n".join(parts), "docx_text"
    except Exception as e:
        return "", f"docx_err:{e.__class__.__name__}"

def extract_excel(path: str) -> Tuple[str, str]:
    """新旧エクセルを読み込み、AIが理解しやすいMarkdown表形式に整形する"""
    out = []
    ext = os.path.splitext(path)[1].lower()
    safe_p = get_safe_path(path)
    try:
        if ext in (".xlsx", ".xlsm") and openpyxl:
            wb = openpyxl.load_workbook(safe_p, data_only=True, read_only=True)
            for ws in wb.worksheets[:10]:
                out.append(f"## Sheet: {ws.title}")
                for row in ws.iter_rows(max_row=400, max_col=40, values_only=True):
                    if any(row):
                        out.append("| " + " | ".join([str(c).strip().replace("\n", " ") if c is not None else "" for c in row]) + " |")
                out.append("")
            wb.close()
            return "\n".join(out), "xlsx_md"
        elif ext == ".xls" and xlrd:
            wb = xlrd.open_workbook(safe_p)
            for sheet_idx in range(min(10, wb.nsheets)):
                ws = wb.sheet_by_index(sheet_idx)
                out.append(f"## Sheet: {ws.name}")
                for row_idx in range(min(400, ws.nrows)):
                    row = ws.row_values(row_idx)
                    if any(row):
                        out.append("| " + " | ".join([str(c).strip().replace("\n", " ") if c else "" for c in row]) + " |")
                out.append("")
            return "\n".join(out), "xls_md"
        else:
            return "", "excel_lib_missing"
    except Exception as e:
        return "", f"excel_err:{e.__class__.__name__}"

def extract_xdw(path: str) -> Tuple[str, str]:
    """DocuWorksからテキストを抽出する。
    xdwlib（Pythonバインディング）を優先し、次にxdw2text.exeを試みる。
    コンソールウィンドウは一切表示しない。"""
    global _XDW2TEXT_PATH
    safe_p = get_safe_path(path)

    # 方法1: xdwlib（Python製DocuWorksバインディング）を優先的に試す
    if XDWLIB_AVAILABLE:
        try:
            doc = xdwlib.xdwopen(path)
            texts = [doc[pg].text for pg in range(doc.pages)]
            doc.close()
            result = "\n".join(texts)
            if result.strip():
                return result, "xdw_xdwlib"
        except Exception:
            pass  # 失敗したらxdw2textにフォールバック

    # 方法2: xdw2text.exe を試す
    # 一度見つかったパスをキャッシュ済みなら1回だけ試す（ウィンドウ多発を防止）
    # まだ見つかっていない場合は全候補を順に試す
    candidates_to_try = [_XDW2TEXT_PATH] if _XDW2TEXT_PATH else XDW2TEXT_CANDIDATES

    for cmd in candidates_to_try:
        if not cmd:
            continue
        try:
            result = subprocess.run(
                [cmd, safe_p],
                capture_output=True,
                text=True,
                encoding="cp932",
                errors="ignore",
                timeout=30,
                **_WIN_NO_CONSOLE,   # ← Windowsのコンソールウィンドウを非表示
            )
            if result.returncode == 0:
                _XDW2TEXT_PATH = cmd  # 使えるexeを記憶して次回以降の探索を省略
                if result.stdout.strip():
                    return result.stdout, "xdw_text"
                return "", "xdw_empty_or_protected"  # ツールは動いたがファイルが空
        except FileNotFoundError:
            if cmd == _XDW2TEXT_PATH:
                _XDW2TEXT_PATH = None  # キャッシュが無効になったのでリセット
            continue
        except Exception:
            continue

    # 方法3: xdoc2txt.exe を試す（無料ツール: https://ebstudio.info/home/xdoc2txt.html）
    # DocuWorks Viewer Light をインストールすると DocuWorks Content Filter (iFilter) が
    # 自動インストールされるため、-i オプションで XDW からテキスト抽出できる。
    global _XDOC2TXT_PATH
    xdoc2txt_candidates = [_XDOC2TXT_PATH] if _XDOC2TXT_PATH else XDOC2TXT_CANDIDATES
    for cmd in xdoc2txt_candidates:
        if not cmd:
            continue
        # まず -i (iFilter) オプションで試す → DocuWorks Viewer Light の iFilter を利用
        for args in [[cmd, "-i", safe_p], [cmd, safe_p]]:
            try:
                result = subprocess.run(
                    args,
                    capture_output=True,
                    text=True,
                    encoding="cp932",
                    errors="ignore",
                    timeout=30,
                    **_WIN_NO_CONSOLE,
                )
                if result.returncode == 0 and result.stdout.strip():
                    _XDOC2TXT_PATH = cmd
                    method_name = "xdw_xdoc2txt_ifilter" if "-i" in args else "xdw_xdoc2txt"
                    return result.stdout, method_name
            except FileNotFoundError:
                if cmd == _XDOC2TXT_PATH:
                    _XDOC2TXT_PATH = None
                break  # このcmdは存在しないので次のcmdへ
            except Exception:
                break

    # DocuWorks Viewer Lightがインストール済みの場合でも、
    # テキスト抽出には別途 xdoc2txt.exe が必要（iFilter経由でXDWを読める）
    return "", "xdw_text_extractor_missing"

def split_main_attach(text: str, kws: List[str]) -> Tuple[str, str]:
    lines = text.splitlines()
    cut_idx = -1
    for i, line in enumerate(lines):
        for k in kws:
            if re.match(k, line):
                cut_idx = i
                break
        if cut_idx != -1: break

    if cut_idx > 5:
        main_text = "\n".join(lines[:cut_idx])
        attach_text = "\n".join(lines[cut_idx:])
        return main_text.strip(), attach_text.strip()
    return text.strip(), ""

# ── 文書タイプ自動判別 ──────────────────────────────────

# フォルダ名による判別キーワード
_DOCTYPE_FOLDER_LAW = re.compile(
    r"法令|法律|政令|省令|規則|施行令|施行規則|条例|告示|訓令|条文"
)
_DOCTYPE_FOLDER_MANUAL = re.compile(
    r"マニュアル|手順書|手引き|てびき|ガイド|ガイドライン|要領|要綱|内規|規程|SOP|社内|内部"
)

# 本文内容による法令判別パターン
_LAW_ARTICLE_RE = re.compile(r"第[一二三四五六七八九十百千\d１-９０]+条")
_LAW_NAME_PATTERNS = [
    r"消防法", r"危険物の規制に関する政令", r"危険物の規制に関する規則",
    r"石油コンビナート等災害防止法", r"液化石油ガスの保安の確保及び取引の適正化に関する法律",
    r"高圧ガス保安法", r"火薬類取締法", r"建築基準法",
]


def _detect_doc_type(rel_path: str, text: str) -> str:
    """文書タイプを自動判別する。

    判定優先順位:
      1. フォルダ名に法令系キーワード → "法令"
      2. フォルダ名にマニュアル系キーワード → "マニュアル"
      3. 本文に条文パターン（第○条）が多数 → "法令"
      4. デフォルト → "通知"
    """
    # フォルダ名による判別（最優先）
    folder_parts = rel_path.replace("\\", "/").rsplit("/", 1)
    folder_path = folder_parts[0] if len(folder_parts) > 1 else ""

    if _DOCTYPE_FOLDER_LAW.search(folder_path):
        return "法令"
    if _DOCTYPE_FOLDER_MANUAL.search(folder_path):
        return "マニュアル"

    # 本文内容による判別（フォルダ名が使えない場合）
    # 条文パターンが5回以上出現 → 法令本文の可能性が高い
    target = text[:10000]
    article_hits = len(_LAW_ARTICLE_RE.findall(target))
    if article_hits >= 5:
        # 条文が多数あっても「通知する」等があれば通知
        if re.search(r"通知する|依頼する|連絡する|送付する", target[:3000]):
            return "通知"
        return "法令"

    return "通知"


def convert_japanese_year(text: str) -> str:
    def replacer(match):
        era = match.group(1)
        year_str = match.group(2)
        year = 1 if year_str == "元" else int(year_str)
        if era == "令和": west_year = 2018 + year
        elif era == "平成": west_year = 1988 + year
        elif era == "昭和": west_year = 1925 + year
        else: return match.group(0)
        return f"{match.group(0)}（{west_year}年）"
    return re.sub(r"(令和|平成|昭和)\s*([0-9元]+)\s*年", replacer, text)

# 通知タイトルの典型的な末尾パターン（日本の公文書）
_TITLE_ENDINGS = (
    r"について[（(]?通知[）)]?\s*$", r"について\s*$", r"に関する件\s*$",
    r"に関して\s*$", r"に係る件\s*$", r"の件\s*$",
)
# ヘッダー行の典型パターン（文書番号・日付・宛先・発出者など）
_HEADER_PATTERNS = (
    r"^第\d+号",
    # 「消防危第」「消防予第」等を正しく検出（[危予施立]に危を追加）
    r"^[消総危]防[危予施立]?第",
    # OCR化けで先頭にゴミ文字が付いた文書番号行（例: "ロロ消防危第284号"）
    r"消防[危予施立]?第\s*\d+\s*号",
    r"^\d{4}年", r"^令和|^平成|^昭和",
    # 宛先・受信者（各都道府県・各指定都市・各消防本部 等）
    r"各都道府県|各消防本部|各市町村|各指定都市|各政令市|各中核市",
    r"消防本部長|消防署長殿|知事殿",
    r"殿\s*$", r"御中\s*$",
    # 発出者（消防庁・総務省・東京消防庁 等）
    r"^消防庁|^総務省|^危険物保安室|^予防課",
    r"^東京消防庁|^各消防本部長|^各消防署長",
    r"官印省略",
    # 防災主管課・消防本部 等（OCR文書で宛先がタイトルに誤検出される対策）
    r"防災主管課", r"^消防[本局]部", r"都市消防本部",
    # 事務連絡・通知文書の定型冒頭行
    r"^事務連絡\s*$", r"^写\s*$", r"^別記\s*$",
)

# ── 箇条書き番号で始まる行（タイトルではなく本文の項目） ──
_NUMBERED_ITEM_RE = re.compile(
    r"^[\s　]*(?:"
    r"[１-９][０-９]*[\s　．.\-\)）]|"      # 全角数字で始まる項目（「１ 」「１．」等）
    r"\d+[\s　．.\-\)）]|"                   # 半角数字で始まる項目（「1.」「1 」等）
    r"[①-⑳]|"                               # 丸数字
    r"（[１-９]）|"                           # （１）等
    r"\([1-9]\)"                             # (1) 等
    r")"
)

# 文章の途中（助詞・接続詞・読点）で始まる行はタイトル候補から除外する
_MID_SENTENCE_RE = re.compile(r"^[てしがのにをはもとなかよりでもし、。・ー…「」]")


def _compute_ocr_quality(text: str) -> float:
    """OCRテキストの品質スコアを0.0〜1.0で返す。
    高い = 良質なテキスト、低い = ゴミが多い。
    テキストPDF・Word・Excel等はデフォルト1.0を使い、この関数はOCR結果のみに適用する。"""
    if not text or not text.strip():
        return 0.0
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if not lines:
        return 0.0
    total_chars = sum(len(l) for l in lines)
    if total_chars == 0:
        return 0.0

    # (1) 日本語文字比率（高い方が良い）
    jp_chars = len(re.findall(r'[ぁ-んァ-ン一-龥]', text))
    jp_ratio = jp_chars / total_chars

    # (2) ゴミ行比率（低い方が良い）
    garbage_count = sum(1 for l in lines if _is_garbage_line(l))
    garbage_ratio = garbage_count / len(lines)

    # (3) 意味のある単語を含む行の比率（高い方が良い）
    # 「について」「に関する」「消防」「危険物」等の通知キーワードで判定
    _meaningful_re = re.compile(
        r"について|に関する|通知|消防|危険物|規則|政令|省令|条例|届出|許可|検査|安全"
    )
    meaningful_lines = sum(1 for l in lines if _meaningful_re.search(l))
    meaningful_ratio = meaningful_lines / len(lines)

    # (4) 平均行長（極端に短い行が多い = OCR断片化）
    avg_len = total_chars / len(lines)
    len_score = min(1.0, avg_len / 25.0)

    # 総合スコア
    score = (jp_ratio * 0.35
             + (1.0 - garbage_ratio) * 0.25
             + meaningful_ratio * 0.20
             + len_score * 0.20)
    return round(min(1.0, max(0.0, score)), 2)


def _is_ocr_garbled_title(s: str) -> bool:
    """OCR由来の壊れたタイトル候補を拒否する。
    例: "河顧客に自ら...", "*品としての特月 8日付け..."
    """
    if not s:
        return True
    # 先頭1〜2文字がランダムな非日本語文字（OCRゴミの典型）
    if re.match(r'^[A-Za-z\*\#\$\@\!\?\~\^\&\%\+\=\|\\\/<>]{1,2}[ぁ-んァ-ン一-龥]', s):
        return True
    # 先頭が孤立した1文字の漢字/カナ + 残りの文脈と不整合
    # 例: "河顧客に..." → "河" は前の行からの誤結合
    if (len(s) >= 10
            and re.match(r'^[ぁ-んァ-ン一-龥]{1}[ぁ-んァ-ン一-龥]', s)
            and s[0] not in 'のはがをにでもとやへ各本全新旧上下前後'):
        # 2文字目以降で明確なタイトルパターンが始まるか確認
        rest = s[1:]
        for pat in _TITLE_ENDINGS:
            if re.search(pat, rest):
                # 先頭1文字を除いてタイトルとして成立 → 先頭はOCRゴミ
                return True
    # 120文字超はタイトルとしては異常に長い（OCRの行結合エラーの可能性大）
    if len(s) > 120:
        return True
    # 途中にOCR化けの典型パターン（ランダムな半角英字が日本語文中に混入）
    # 例: "Sいて、可搬式の" → "S" は "さ" のOCR化け
    fragments = re.findall(r'[A-Z][ぁ-んァ-ン一-龥]', s)
    if len(fragments) >= 2:
        return True
    return False


# ── 改廃関係の検出パターン ──
_AMENDMENT_RE = re.compile(
    r"(「[^」]{3,60}」\s*(?:を|の)\s*(?:一部改正|全部改正|廃止|制定|追加|削除))"
    r"|((?:一部|全部)?(?:改正|廃止)(?:する|した|され))"
    r"|(新たに(?:制定|公布|施行))"
)

# ── 関連法令番号の抽出パターン ──
_LAW_REF_RE = re.compile(
    r"(?:危険物の規制に関する)?(?:政令|規則|省令|法律|法|条例|告示|訓令)"
    r"(?:\s*第\s*\d+\s*条(?:\s*の\s*\d+)?(?:\s*第\s*\d+\s*項)?(?:\s*第\s*\d+\s*号)?)?"
)


def _extract_related_laws(text: str) -> List[str]:
    """テキストから関連法令の参照（「政令第○条」等）を抽出する"""
    target = text[:6000]
    hits = _LAW_REF_RE.findall(target)
    # 重複除去して返す（出現順を維持）
    seen = set()
    result = []
    for h in hits:
        h = h.strip()
        if h and len(h) >= 4 and h not in seen:
            seen.add(h)
            result.append(h)
    return result[:10]  # 最大10件


def _extract_amendments(text: str) -> List[str]:
    """テキストから改廃関係の情報を抽出する"""
    target = text[:6000]
    hits = _AMENDMENT_RE.findall(target)
    result = []
    for groups in hits:
        for g in groups:
            g = g.strip()
            if g and len(g) >= 4 and g not in result:
                result.append(g)
    return result[:5]  # 最大5件


def _date_to_sort_key(date_str: str) -> str:
    """日付文字列をYYYYMMDD形式のソートキーに変換する"""
    if not date_str:
        return "99999999"
    # 西暦表記（「2023年3月1日」）
    m = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', date_str)
    if m:
        return f"{m.group(1)}{int(m.group(2)):02d}{int(m.group(3)):02d}"
    # 和暦のカッコ内西暦（「令和5年（2023年）」等 — convert_japanese_yearで追加）
    m = re.search(r'（(\d{4})年）', date_str)
    if m:
        # 月日も取る
        md = re.search(r'(\d{1,2})\s*月\s*(\d{1,2})\s*日', date_str)
        if md:
            return f"{m.group(1)}{int(md.group(1)):02d}{int(md.group(2)):02d}"
        return f"{m.group(1)}0101"
    return "99999999"


def _is_meaningful_title(s: str) -> bool:
    """
    タイトルとして意味のある文字列かを判定する。
    ・日本語文字（ひらがな・カタカナ・漢字）が1文字以上あること
    ・全体の15%以上が日本語文字であること
    → OCRゴミ行（"NMWMMMMMUMNMNI" 等）を拒否する
    """
    if not s:
        return False
    jp_count = len(re.findall(r'[ぁ-んァ-ン一-龥]', s))
    if jp_count == 0:
        return False
    return jp_count / len(s) >= 0.15


def _is_similar_to_title(line: str, title: str) -> bool:
    """概要の行がタイトルと内容的に重複しているかを判定する。
    概要冒頭にタイトルがそのまま繰り返されるのを防止するために使う。"""
    if not title or len(line) < 6:
        return False
    # 完全一致・包含関係
    if line in title or title in line:
        return True
    # 空白・句読点を除去して比較
    _strip_re = re.compile(r'[\s　、。・（）\(\)\-\—\―]')
    clean_line = _strip_re.sub('', line)
    clean_title = _strip_re.sub('', title)
    if clean_title and clean_line:
        if clean_line in clean_title or clean_title in clean_line:
            return True
    return False


def guess_title(text: str, fallback: str) -> str:
    """通知タイトルを推定する。
    「〜について」パターンを優先し、OCRゴミ・箇条書き番号・ヘッダー行を厳密に拒否する。"""
    lines = text.splitlines()

    def _is_title_connectable(line_text: str) -> bool:
        """前行・前々行がタイトルの一部として結合可能かを判定する"""
        return (5 <= len(line_text) <= 120
                and not any(re.search(p, line_text) for p in _HEADER_PATTERNS)
                and not _MID_SENTENCE_RE.match(line_text)
                and not _NUMBERED_ITEM_RE.match(line_text)
                and _is_meaningful_title(line_text)
                and not _is_ocr_garbled_title(line_text)
                and not any(re.search(pat, line_text) for pat in _TITLE_ENDINGS))

    def _validate_title(candidate: str) -> Optional[str]:
        """タイトル候補の最終バリデーション（OCRゴミ・異常長を拒否）"""
        if not candidate or len(candidate) > 120:
            return None
        if _is_ocr_garbled_title(candidate):
            return None
        if _NUMBERED_ITEM_RE.match(candidate):
            return None
        if not _is_meaningful_title(candidate):
            return None
        return candidate

    # パターン1: 「〜について」「〜に関する件」で終わる行を優先（通知タイトルの典型形）
    # 複数行（最大3行）にまたがるタイトルにも対応
    for i, line in enumerate(lines[:100]):
        s = line.strip()

        # タイトル末尾パターンに一致する行（10文字以上、120文字以内）
        if 10 <= len(s) <= 120 and any(re.search(pat, s) for pat in _TITLE_ENDINGS):
            # OCRゴミチェック
            if _is_ocr_garbled_title(s):
                continue
            # 箇条書き番号で始まる行はタイトルではない
            if _NUMBERED_ITEM_RE.match(s):
                continue
            # 前行がヘッダーでなく意味のある行なら結合してタイトルを補完
            if i > 0:
                prev = lines[i - 1].strip()
                if _is_title_connectable(prev):
                    # さらに前々行も結合可能か確認（3行にまたがるタイトル）
                    if i > 1:
                        prev2 = lines[i - 2].strip()
                        if _is_title_connectable(prev2):
                            result = _validate_title(prev2 + prev + s)
                            if result:
                                return result
                    result = _validate_title(prev + s)
                    if result:
                        return result
            return s

        # タイトル末尾パターンに一致するが短い行（< 10文字）→ 前行と結合
        if 3 <= len(s) <= 9 and any(re.search(pat, s) for pat in _TITLE_ENDINGS):
            if i > 0:
                prev = lines[i - 1].strip()
                if _is_title_connectable(prev):
                    if i > 1:
                        prev2 = lines[i - 2].strip()
                        if _is_title_connectable(prev2):
                            result = _validate_title(prev2 + prev + s)
                            if result:
                                return result
                    result = _validate_title(prev + s)
                    if result:
                        return result

        # 短い行が続いて次行でタイトルが完結するケース
        if 3 <= len(s) < 10 and i + 1 < len(lines):
            next_s = lines[i + 1].strip()
            combined = s + next_s
            if 10 <= len(combined) <= 120 and any(re.search(pat, combined) for pat in _TITLE_ENDINGS):
                result = _validate_title(combined)
                if result:
                    return result

    # パターン2: ヘッダー行・文中断片をスキップして最初の意味のある行を取る
    for li, line in enumerate(lines[:80]):
        s = line.strip()
        if len(s) < 8 or len(s) > 120:
            continue
        if re.match(r"^[\d\-\s\(\)（）・ 　]+$", s):
            continue
        if any(re.search(p, s) for p in _HEADER_PATTERNS):
            continue
        if _MID_SENTENCE_RE.match(s):
            continue
        if not _is_meaningful_title(s):
            continue
        if _is_ocr_garbled_title(s):
            continue
        if _NUMBERED_ITEM_RE.match(s):
            continue
        # 次行と結合するとタイトルになる場合は結合版を返す
        if li + 1 < len(lines):
            next_s = lines[li + 1].strip()
            combined = s + next_s
            result = _validate_title(combined)
            if result and any(re.search(pat, combined) for pat in _TITLE_ENDINGS):
                return result
        return s
    return fallback


def guess_title_law(text: str, fallback: str) -> str:
    """法令文書のタイトルを推定する。
    法令名（「消防法」「危険物の規制に関する政令」等）を検出する。"""
    lines = text.splitlines()

    # パターン1: 既知の法令名を直接検出
    known_laws = [
        "危険物の規制に関する政令", "危険物の規制に関する規則",
        "消防法施行令", "消防法施行規則", "消防法",
        "石油コンビナート等災害防止法", "高圧ガス保安法",
        "液化石油ガスの保安の確保及び取引の適正化に関する法律",
        "火薬類取締法", "建築基準法",
    ]
    for i, line in enumerate(lines[:30]):
        s = line.strip()
        for law_name in known_laws:
            if law_name in s and len(s) <= 80:
                # 「〜の一部を改正する〜」のようなタイトルも拾う
                if "改正" in s and len(s) >= 10:
                    return s
                return law_name

    # パターン2: 「第一章 総則」等の章立てがある → その前に法令名がある
    for i, line in enumerate(lines[:50]):
        s = line.strip()
        if re.match(r"^第[一二三四五六七八九十]+章", s):
            # この行より前で最後の意味のある行が法令名
            for j in range(i - 1, -1, -1):
                prev = lines[j].strip()
                if prev and len(prev) >= 4 and len(prev) <= 80:
                    if not re.match(r"^[\d\s（）\(\)]+$", prev):
                        return prev
            break

    # パターン3: 先頭の意味のある行を取る（法令ファイルなのでヘッダーパターンは適用しない）
    for line in lines[:20]:
        s = line.strip()
        if not s or len(s) < 4 or len(s) > 80:
            continue
        if re.match(r"^[\d\s\-（）\(\)・ 　]+$", s):
            continue
        if _is_garbage_line(s):
            continue
        return s

    return fallback


def guess_title_manual(text: str, fallback: str) -> str:
    """マニュアル・手順書のタイトルを推定する。"""
    lines = text.splitlines()

    # 先頭付近から最初の意味のある行を取る（ヘッダーパターンは適用しない）
    for line in lines[:30]:
        s = line.strip()
        if not s or len(s) < 4 or len(s) > 120:
            continue
        if re.match(r"^[\d\s\-（）\(\)・ 　]+$", s):
            continue
        if _is_garbage_line(s):
            continue
        return s
    return fallback


def guess_date(text: str) -> str:
    m = re.search(r"(令和|平成|昭和)\s*[0-9元]+\s*年\s*\d+\s*月\s*\d+\s*日(（\d{4}年）)?", text)
    if m: return m.group(0)
    m2 = re.search(r"\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日", text)
    return m2.group(0) if m2 else ""

def guess_issuer(text: str) -> str:
    for cand in ["消防庁", "総務省消防庁", "消防局", "危険物保安室", "予防課"]:
        if cand in text: return cand
    return ""

def tag_text(text: str) -> Tuple[List[str], List[str], Dict[str, List[str]]]:
    ev: Dict[str, List[str]] = {}; fac: List[str] = []; work: List[str] = []
    target = text[:8000]
    for t, ps in FACILITY_TAGS.items():
        if hits := [p for p in ps if re.search(p, target)]:
            fac.append(t); ev[t] = hits[:3]
    for t, ps in WORK_TAGS.items():
        if hits := [p for p in ps if re.search(p, target)]:
            work.append(t); ev[t] = hits[:3]
    # ※「共通」フォールバックは廃止。施設が特定できない通知はタグなしとする。
    return fac, work, ev

def _normalize_line(s: str) -> str:
    """PDF抽出由来の行内スペースを正規化する"""
    # 日本語文字間の不要スペースを除去（例: "令 和 3 年" → "令和3年"）
    # ※ 数字↔日本語間のスペースは箇条書き番号等で意味があるので除去しない
    s = re.sub(r'([ぁ-んァ-ン一-龥])[ \t]+([ぁ-んァ-ン一-龥])', r'\1\2', s)
    # 連続する半角スペースを1つに（全角スペース・先頭インデントは保持）
    s = re.sub(r'[ \t]{2,}', ' ', s)
    return s


# ── 箇条書き番号の先頭パターン（階層構造を保持するために使う） ──
# 例: 「１」「（１）」「ア」「(ア)」「①」「・」「(1)」「1.」等
_BULLET_RE = re.compile(
    r"^[\s　]*(?:"
    r"[①-⑳]|"                          # 丸数字 ①②…
    r"[１-９０][０-９]*[．.、]|"          # 全角数字+句点
    r"（[１-９０ア-ン一-龥]{1,3}）|"     # （全角）
    r"\([1-9ア-ンa-z]{1,3}\)|"          # (半角)
    r"[ア-ン][．.、\s]|"                 # カタカナ+句点
    r"[a-zA-Z][．.、\s]|"               # アルファベット
    r"[・•◆▶▷]\s"                      # 中黒・記号
    r")"
)

# ── OCRゴミ行のパターン（校閲官合意：除去対象） ──
_GARBAGE_LINE_RE = re.compile(
    r"^[\s　]*(?:"
    r"[■□◆◇▲△▼▽●○◎★☆※＊〒]{1,3}|"   # 記号のみ行
    r"[-─━=＝]{4,}|"                     # 罫線のみ行
    r"[・\-]{3,}"                         # 点線・ハイフン3個以上
    r")[\s　]*$"
)

# ── 終端行パターン ──
_TERMINATOR_RE = re.compile(
    r"^\s*(以上|以下余白|（了）|－\s*了\s*－|【以上】|〔以上〕)\s*$"
)

# ── フッター行パターン（担当者・問い合わせ先・電話番号・ページ番号等） ──
_FOOTER_LINE_RE = re.compile(
    r"担当(?:係)?[:：]|問い?合わせ(?:先)?[:：]|電話[:：]|FAX[:：\s]|℡|"
    # TEL（大文字・小文字・混合）に続く数字、または単独で行に存在する場合
    r"TEL[:：\s]*\d|Tel[:：\s]*\d|tel[:：\s]*\d|"
    # 電話番号の形式（00-0000-0000、0120-000-000 等）
    r"\d{2,5}[-－]\d{3,4}[-－]\d{4}|"
    r"〒\d{3}[-－]?\d{4}|問合せ|照会先|内線|直通|"
    r"^\s*[－\-]\s*\d+\s*[－\-]\s*$"
)

# ── 趣旨文の末尾パターン（日本の公文書定型：「通知する。」等で終わる行） ──
_INTENT_SENTENCE_END_RE = re.compile(
    r"(通知する|通知します|伝達する|送付する|連絡する|回答する|依頼する"
    r"|お知らせする|お伝えする|送付します|依頼します)。?\s*$"
)

# ── 施行日・適用日のパターン ──
_ENFORCEMENT_DATE_RE = re.compile(
    r"(?:施行|適用|公布|発効|実施|以降|より)"
    r".{0,6}"
    r"(?:令和|平成|昭和)\s*[0-9元]+\s*年\s*\d+\s*月\s*\d+\s*日"
    r"|"
    r"(?:令和|平成|昭和)\s*[0-9元]+\s*年\s*\d+\s*月\s*\d+\s*日"
    r".{0,10}(?:施行|適用|公布|発効|以降|から)"
)


def _is_garbage_line(s: str) -> bool:
    """OCRゴミ・孤立記号・罫線・日本語皆無行などの除去すべき行か判定する"""
    if not s:
        return False
    # 1〜2文字のみ（記号・数字・カナ等）は除去
    if len(s) <= 2 and re.match(r"^[^\u3041-\u9FFF]*$", s):
        return True
    if _GARBAGE_LINE_RE.match(s):
        return True
    # OCRゴミ検出: スペースを除いた文字で判定
    no_space = s.replace(' ', '').replace('　', '').replace('\t', '')
    if len(no_space) >= 4:
        jp_count = len(re.findall(r'[ぁ-んァ-ン一-龥]', no_space))
        total = len(no_space)
        # (1) 日本語文字が一切ない → OCRゴミ
        if jp_count == 0 and total >= 6:
            return True
        # (2) 日本語比率が極端に低い（10%未満で10文字以上）→ OCR化け
        #     例: "MNWMれMMNI" のようなケース
        if total >= 10 and jp_count > 0 and (jp_count / total) < 0.10:
            return True
        # (3) 連続するASCII大文字が多い → OCR化けの典型
        #     例: "NMWMMMMMUMNMNI" の中にカタカナ1文字混入
        ascii_upper_runs = re.findall(r'[A-Z]{4,}', no_space)
        if ascii_upper_runs and sum(len(r) for r in ascii_upper_runs) > total * 0.5:
            return True
    return False


def _is_header_or_footer(s: str) -> bool:
    """ヘッダー（発出者・宛先・文書番号）またはフッター行か判定する"""
    return bool(
        any(re.search(p, s) for p in _HEADER_PATTERNS)
        or _FOOTER_LINE_RE.search(s)
    )


def _join_short_continuation_lines(lines: List[str]) -> List[str]:
    """
    PDF抽出で途切れた短い行を次行と連結する。
    ─ 校閲官合意ルール ─
    ・ゴミ行・終端行は連結しない（そのまま渡して後段でフィルタ）
    ・行末が句読点「。」「、」で終わっている → 完結行なので連結しない
    ・行頭が箇条書き番号 → 新項目の開始なので連結しない
    ・行の長さが10文字未満かつ上記に該当しない → 次行の先頭に連結
    """
    result: List[str] = []
    i = 0
    while i < len(lines):
        s = lines[i]
        # ゴミ行・終端行はそのまま（次行と混ぜない）
        if _is_garbage_line(s) or _TERMINATOR_RE.match(s):
            result.append(s)
            i += 1
            continue
        # 短い行で、次行があり、箇条書き番号で始まらず、句点で終わらない → 連結
        if (len(s) < 10
                and i + 1 < len(lines)
                and not _BULLET_RE.match(s)
                and not _is_garbage_line(lines[i + 1])
                and not re.search(r"[。、」）\)]\s*$", s)):
            result.append(s + lines[i + 1])
            i += 2
            continue
        result.append(s)
        i += 1
    return result


def _extract_enforcement_date(text: str) -> str:
    """テキストから施行日・適用日を抽出して整形文字列を返す"""
    m = _ENFORCEMENT_DATE_RE.search(text)
    if m:
        # 年月日部分だけ取り出す
        date_m = re.search(
            r"(?:令和|平成|昭和)\s*[0-9元]+\s*年\s*\d+\s*月\s*\d+\s*日",
            m.group(0)
        )
        if date_m:
            return date_m.group(0)
    return ""


def _format_summary(core: str, n: int, title_hint: str = "") -> str:
    """
    概要テキストを読みやすく整形する。

    ─ 専門家会議の合意設計 ─
    1. ゴミ行・ヘッダー行・フッター行を除去
    2. 冒頭のタイトル行をスキップ（タイトル欄との重複を防止）
    3. 短い途切れ行を次行と連結（箇条書き行は除外）
    4. 連続空行を1つに間引く
    5. 終端行（以上・了等）でストップ
    6. 文字数上限でカット
    """
    # 前処理: 行ごとにスペース正規化
    raw_lines = [_normalize_line(l.strip()) for l in core.splitlines()]
    # 短い途切れ行を連結
    merged = _join_short_continuation_lines(raw_lines)

    result_lines: List[str] = []
    char_count = 0
    prev_blank = False
    # 冒頭フェーズ: まだ本文が始まっていない段階でタイトル行をスキップ
    initial_phase = True

    for line in merged:
        stripped = line.strip()

        # 終端行でストップ
        if _TERMINATOR_RE.match(stripped):
            break

        # 空行処理（連続空行を1つに）
        if not stripped:
            if result_lines and not prev_blank:
                result_lines.append("")
            prev_blank = True
            continue
        prev_blank = False

        # OCRゴミ行・ヘッダー・フッターをスキップ
        if _is_garbage_line(stripped):
            continue
        if _is_header_or_footer(stripped):
            continue

        # 冒頭フェーズ: タイトル行が概要に重複表示されるのを防止
        if initial_phase:
            # タイトル末尾パターン（「〜について」等）に一致する行はスキップ
            if any(re.search(pat, stripped) for pat in _TITLE_ENDINGS) and len(stripped) <= 200:
                continue
            # title_hintと内容が重複する行をスキップ
            if title_hint and _is_similar_to_title(stripped, title_hint):
                continue
            # タイトルでもヘッダーでもない最初の実質行 → 本文開始
            initial_phase = False

        result_lines.append(stripped)
        char_count += len(stripped)
        if char_count >= n:
            break

    result = "\n".join(result_lines).rstrip()
    return result[:n] + ("…" if len(result) > n else "")


def make_summary(main_text: str, n: int, title_hint: str = "",
                 ocr_quality: float = 1.0) -> str:
    """
    危険物行政通知の概要を生成する。

    ─ 専門家会議（危険物行政専門家・日本語校閲官・消防職員）合意設計 ─

    【出力構造】
      [趣旨] 本文冒頭の目的文（最大2文・150文字以内）
      [要点] 「記」以降の本文（箇条書き番号・階層構造を保持）
      [施行・適用] 施行日・適用日（自動検出時のみ末尾に付記）

    【除去対象】
      ・宛先・発出者・文書番号行（ヘッダー）
      ・担当者・問い合わせ先行（フッター）
      ・タイトル行（タイトル欄と重複するため）
      ・OCRゴミ行（1〜2文字行、記号だけの行）
    """
    if not main_text.strip():
        return ""

    # OCR品質が極めて低い場合は概要を抑制
    if ocr_quality < 0.25:
        return "（OCR品質が低いため概要を自動生成できません。元ファイルを直接ご確認ください。）"

    # ── Step 1: 施行日を先にテキスト全体から抽出 ──
    enforcement_date = _extract_enforcement_date(main_text)

    # ── Step 2: 「記」の有無で分岐 ──
    ki_match = re.search(r"\n\s*記\s*\n", main_text)

    if ki_match:
        # 【記あり】趣旨（記より前）+ 記以降の要点
        pre_ki  = main_text[:ki_match.start()]
        post_ki = main_text[ki_match.end():]

        # 趣旨: 「〜通知する。」等の趣旨文を1〜2文だけ取る。
        intent_buf = ""
        intent_result = ""
        for raw in pre_ki.splitlines():
            s = _normalize_line(raw.strip())
            if not s or _is_garbage_line(s) or _is_header_or_footer(s):
                continue
            intent_buf += s
            if any(re.search(pat, intent_buf) for pat in _TITLE_ENDINGS):
                intent_buf = ""
                continue
            if title_hint and _is_similar_to_title(intent_buf, title_hint):
                intent_buf = ""
                continue
            if _INTENT_SENTENCE_END_RE.search(intent_buf):
                intent_result = intent_buf
                break
            if re.search(r"。\s*$", intent_buf):
                intent_result = intent_buf
                break
            if len(intent_buf) >= 200:
                intent_result = intent_buf[:200]
                break
        intent_chars = len(intent_result)

        # 要点: 記以降を整形（タイトルヒント付き）
        body_reserve = n - intent_chars - 40  # ラベル分の余裕
        body_part = _format_summary(post_ki, max(200, body_reserve), title_hint=title_hint)

        parts: List[str] = []
        if intent_result:
            parts.append(f"[趣旨] {intent_result}")
        if body_part:
            parts.append(f"[要点]\n{body_part}")
        combined = "\n".join(parts)

    else:
        # 【記なし】タイトル行以降の本文を使う
        lines = main_text.splitlines()
        start = 0

        # タイトル行（「〜について」等）を探してその次行から開始
        for i, line in enumerate(lines[:80]):
            s = line.strip()
            if re.search(r"について|に関する|に関して|に係る", s) and 10 <= len(s) <= 200:
                start = i + 1
                break
            if title_hint and _is_similar_to_title(s, title_hint) and len(s) >= 8:
                start = i + 1
                break

        # フォールバック: 最初の意味のある非ヘッダー行をタイトルとみなす
        if start == 0:
            for i, line in enumerate(lines[:80]):
                s = line.strip()
                if not s or len(s) < 8 or len(s) > 150:
                    continue
                if any(re.search(p, s) for p in _HEADER_PATTERNS):
                    continue
                if _MID_SENTENCE_RE.match(s):
                    continue
                if not _is_meaningful_title(s):
                    continue
                start = i + 1
                break

        # タイトル直後のヘッダー行をスキップ
        skip_end = min(len(lines), start + 15)
        while start < skip_end:
            s = lines[start].strip() if start < len(lines) else ""
            if not s or len(s) < 5 or _is_header_or_footer(s):
                start += 1
            else:
                break

        body_text = "\n".join(lines[start:])
        body_formatted = _format_summary(body_text, n - 20, title_hint=title_hint)

        # 趣旨文を本文先頭から抽出（句点で終わる最初の文）
        intent_part = ""
        rest_part = body_formatted
        for bline in body_formatted.splitlines():
            if re.search(r"。\s*$", bline) or _INTENT_SENTENCE_END_RE.search(bline):
                intent_part = bline
                rest_idx = body_formatted.index(bline) + len(bline)
                rest_part = body_formatted[rest_idx:].strip()
                break

        parts: List[str] = []
        if intent_part:
            parts.append(f"[趣旨] {intent_part}")
            if rest_part:
                parts.append(f"[要点]\n{rest_part}")
        else:
            if body_formatted:
                parts.append(body_formatted)
        combined = "\n".join(parts)

    # ── Step 3: 施行日を末尾に付記（未包含の場合のみ） ──
    if enforcement_date and enforcement_date not in combined:
        suffix = f"\n[施行・適用] {enforcement_date}"
        if len(combined) + len(suffix) <= n + 40:
            combined += suffix

    return combined[:n] + ("…" if len(combined) > n else "")


def make_summary_law(text: str, n: int, title_hint: str = "") -> str:
    """法令文書の概要を生成する。
    条文構造（第○条）を認識し、目的条項・主要条文を抽出する。"""
    if not text.strip():
        return ""

    lines = text.splitlines()
    parts: List[str] = []

    # ── 目的条項を探す（第1条 or 第一条） ──
    purpose_text = ""
    for i, line in enumerate(lines):
        s = line.strip()
        if re.match(r"^第[一1１]条", s):
            # 目的条の内容を収集（次の条文まで）
            buf = [s]
            for j in range(i + 1, min(i + 20, len(lines))):
                next_s = lines[j].strip()
                if re.match(r"^第[二三四五2-9２-９]", next_s):
                    break
                if next_s:
                    buf.append(next_s)
            purpose_text = "".join(buf)
            if len(purpose_text) > 300:
                purpose_text = purpose_text[:300] + "…"
            break

    if purpose_text:
        parts.append(f"[目的] {purpose_text}")

    # ── 章立て構造を抽出 ──
    chapters = []
    for line in lines:
        s = line.strip()
        m = re.match(r"^(第[一二三四五六七八九十百]+章)\s*(.*)", s)
        if m:
            chapters.append(f"{m.group(1)} {m.group(2)}")
    if chapters:
        parts.append("[構成]\n" + "\n".join(chapters[:15]))

    # ── 主要条文の見出しを抽出 ──
    article_heads = []
    for line in lines:
        s = line.strip()
        m = re.match(r"^(第[一二三四五六七八九十百千\d１-９０]+条(?:の[一二三四五六七八九十\d１-９０]+)?)\s*[（(]([^）)]+)[）)]", s)
        if m:
            article_heads.append(f"{m.group(1)}（{m.group(2)}）")
    if article_heads and not chapters:
        parts.append("[条文構成]\n" + "\n".join(article_heads[:20]))

    # ── 施行日 ──
    enforcement_date = _extract_enforcement_date(text)
    if enforcement_date:
        parts.append(f"[施行日] {enforcement_date}")

    if not parts:
        # フォールバック: 先頭の意味のある内容を返す
        return _format_summary(text, n, title_hint=title_hint)

    combined = "\n".join(parts)
    return combined[:n] + ("…" if len(combined) > n else "")


def make_summary_manual(text: str, n: int, title_hint: str = "") -> str:
    """マニュアル・手順書の概要を生成する。
    構造を壊さず、冒頭の目的・対象・手順の要約を抽出する。"""
    if not text.strip():
        return ""

    lines = text.splitlines()
    parts: List[str] = []

    # ── 目的・趣旨を探す ──
    purpose_text = ""
    for i, line in enumerate(lines[:100]):
        s = line.strip()
        if re.search(r"目的|趣旨|はじめに|概要|対象", s) and len(s) >= 4:
            # この行以降の内容を収集
            buf = []
            for j in range(i + 1, min(i + 10, len(lines))):
                next_s = lines[j].strip()
                if not next_s:
                    if buf:
                        break
                    continue
                buf.append(next_s)
            if buf:
                purpose_text = "".join(buf)
                if len(purpose_text) > 200:
                    purpose_text = purpose_text[:200] + "…"
                break

    if purpose_text:
        parts.append(f"[目的] {purpose_text}")

    # ── 見出し構造を抽出 ──
    headings = []
    for line in lines[:200]:
        s = line.strip()
        # 番号付き見出し（「1. 」「第1章」「(1)」等）
        if re.match(r"^(?:\d+[\.．\s]|第\d+[章節項]|[（(]\d+[）)])", s) and 5 <= len(s) <= 80:
            headings.append(s)
    if headings:
        parts.append("[構成]\n" + "\n".join(headings[:15]))

    if not parts:
        # フォールバック: 冒頭の内容をそのまま返す
        result_lines = []
        char_count = 0
        for line in lines:
            s = line.strip()
            if not s or _is_garbage_line(s):
                continue
            result_lines.append(s)
            char_count += len(s)
            if char_count >= n:
                break
        combined = "\n".join(result_lines)
        return combined[:n] + ("…" if len(combined) > n else "")

    combined = "\n".join(parts)
    return combined[:n] + ("…" if len(combined) > n else "")


_ILLEGAL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")

def _xls_safe(s) -> str:
    """Excelに書き込めない制御文字を除去する"""
    if not isinstance(s, str):
        return s
    return _ILLEGAL_CHARS_RE.sub("", s)

def write_excel_index(outdir: str, records: List[Record]):
    if not openpyxl: return

    # ── 色定義 ──────────────────────────────────────────────────
    HEADER_BG   = PatternFill(fill_type="solid", fgColor="1E3A8A")   # 濃青
    OK_BG       = PatternFill(fill_type="solid", fgColor="DCFCE7")   # 薄緑
    REV_BG      = PatternFill(fill_type="solid", fgColor="FEE2E2")   # 薄赤
    HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
    WRAP_CENTER = Alignment(horizontal="center", vertical="top", wrap_text=True)
    WRAP_LEFT   = Alignment(horizontal="left",   vertical="top", wrap_text=True)

    wb = openpyxl.Workbook()

    # ── シート①: 文書一覧 ──────────────────────────────────────
    ws = wb.active
    ws.title = "文書一覧"

    headers = ["No.", "タイプ", "タイトル(推定)", "日付(推定)", "発出者", "施設タグ", "業務タグ", "状態", "理由", "概要", "元ファイル"]
    ws.append(headers)

    # ヘッダー行の書式
    for col_idx in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill   = HEADER_BG
        cell.font   = HEADER_FONT
        cell.alignment = WRAP_CENTER
    ws.row_dimensions[1].height = 30

    # データ行
    for seq, r in enumerate(records, start=1):
        status = "要確認" if r.needs_review else "正常"
        summary_short = _xls_safe(r.summary[:400] if r.summary else "")
        ws.append([
            seq,
            r.doc_type,
            _xls_safe(r.title_guess),
            _xls_safe(r.date_guess),
            _xls_safe(r.issuer_guess),
            " / ".join(r.tags_facility),
            " / ".join(r.tags_work),
            status,
            _xls_safe(r.reason),
            summary_short,
            _xls_safe(r.relpath),
        ])
        row_num = seq + 1
        fill = REV_BG if r.needs_review else OK_BG
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=row_num, column=col_idx)
            cell.fill = fill
            cell.alignment = WRAP_LEFT
        # タイプ列・状態列はセンタリング
        ws.cell(row=row_num, column=2).alignment = WRAP_CENTER
        ws.cell(row=row_num, column=8).alignment = WRAP_CENTER
        # 「要確認」セルは赤字で強調
        if r.needs_review:
            ws.cell(row=row_num, column=8).font = Font(bold=True, color="DC2626")

    # 列幅（近似値）
    col_widths = [6, 10, 42, 20, 14, 24, 24, 8, 32, 55, 50]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # フリーズとオートフィルター
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions

    # ── シート②: サマリー ──────────────────────────────────────
    ws2 = wb.create_sheet("サマリー")
    ok_count  = sum(1 for r in records if not r.needs_review)
    rev_count = len(records) - ok_count

    def _s2_header(row, label):
        cell = ws2.cell(row=row, column=1, value=label)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = HEADER_BG
        cell.alignment = WRAP_CENTER

    ws2.append(["集計項目", "件数"])
    _s2_header(1, "集計項目")
    ws2.cell(row=1, column=2).font   = HEADER_FONT
    ws2.cell(row=1, column=2).fill   = HEADER_BG
    ws2.cell(row=1, column=2).alignment = WRAP_CENTER

    ws2.append(["総ファイル数", len(records)])
    ws2.append(["正常抽出",     ok_count])
    ws2.append(["要確認",       rev_count])
    ws2.append([""])

    ws2.append(["施設タグ別件数", ""])
    _s2_header(ws2.max_row, "施設タグ別件数")
    tag_fac: Dict[str, int] = {}
    for r in records:
        for t in r.tags_facility:
            tag_fac[t] = tag_fac.get(t, 0) + 1
    for t, c in sorted(tag_fac.items(), key=lambda x: -x[1]):
        ws2.append([t, c])

    ws2.append([""])
    ws2.append(["業務タグ別件数", ""])
    _s2_header(ws2.max_row, "業務タグ別件数")
    tag_work: Dict[str, int] = {}
    for r in records:
        for t in r.tags_work:
            tag_work[t] = tag_work.get(t, 0) + 1
    for t, c in sorted(tag_work.items(), key=lambda x: -x[1]):
        ws2.append([t, c])

    ws2.append([""])
    ws2.append(["要確認の理由別", ""])
    _s2_header(ws2.max_row, "要確認の理由別")
    reason_counts: Dict[str, int] = {}
    for r in records:
        if r.needs_review and r.reason:
            reason_counts[r.reason] = reason_counts.get(r.reason, 0) + 1
    for reason, cnt in sorted(reason_counts.items(), key=lambda x: -x[1]):
        ws2.append([reason, cnt])

    ws2.column_dimensions["A"].width = 50
    ws2.column_dimensions["B"].width = 10

    excel_path = os.path.join(outdir, "00_統合目次.xlsx")
    try:
        wb.save(excel_path)
    except PermissionError:
        raise PermissionError("00_統合目次.xlsx が他のアプリで開かれています。閉じてからやり直してください。")

def write_md_indices(outdir: str, records: List[Record]):
    with open(os.path.join(outdir, "00_統合目次.md"), "w", encoding="utf-8") as f:
        f.write("# 統合目次（法令・通知・マニュアル）\n\n")
        current_type = ""
        for r in records:
            # タイプが変わったらセクション見出しを出力
            if r.doc_type != current_type:
                current_type = r.doc_type
                type_counts = sum(1 for x in records if x.doc_type == current_type)
                f.write(f"## {current_type}（{type_counts}件）\n\n")

            laws_str = f"\n  - 関連法令: {', '.join(r.related_laws)}" if r.related_laws else ""
            amend_str = f"\n  - 改廃: {', '.join(r.amendments)}" if r.amendments else ""
            ocr_str = f"\n  - OCR品質: {r.ocr_quality:.0%}" if r.ocr_quality < 1.0 else ""
            f.write(
                f"- **[{r.doc_type}] {r.title_guess}**\n"
                f"  - 日付: {r.date_guess} / 発出: {r.issuer_guess}\n"
                f"  - タグ: [{'/'.join(r.tags_facility)}] [{'/'.join(r.tags_work)}]"
                f"{laws_str}{amend_str}{ocr_str}\n"
                f"  - 概要: {r.summary}\n"
                f"  - 元: `{r.relpath}`\n\n"
            )

def write_binded_texts(outdir: str, records: List[Record], limit_bytes: int):
    """文書タイプ別にNotebookLM用テキストを出力する。
    法令→通知→マニュアルの順に、タイプ別ファイル名で出力。"""

    # タイプ別にグループ化（出力順: 法令 → 通知 → マニュアル）
    type_order = {"法令": 1, "通知": 2, "マニュアル": 3}
    type_groups: Dict[str, List[Record]] = {"法令": [], "通知": [], "マニュアル": []}
    for r in records:
        group = r.doc_type if r.doc_type in type_groups else "通知"
        type_groups[group].append(r)

    type_prefixes = {"法令": "01_法令", "通知": "02_通知", "マニュアル": "03_マニュアル"}

    for doc_type in ["法令", "通知", "マニュアル"]:
        group_records = type_groups[doc_type]
        if not group_records:
            continue

        prefix = type_prefixes[doc_type]
        chunk_idx = 1
        current_size = 0
        current_blocks: List[str] = []
        doc_num = 0

        def flush(p=prefix, ci=[chunk_idx], cs=[current_size], cb=current_blocks):
            if not cb:
                return
            fname = f"NotebookLM用_{p}_{ci[0]:02d}.txt"
            with open(os.path.join(outdir, fname), "w", encoding="utf-8") as f:
                f.write("\n".join(cb))
            ci[0] += 1
            cs[0] = 0
            cb.clear()

        for r in group_records:
            if not r.full_text_for_bind.strip():
                continue
            doc_num += 1

            block = (
                f"\n\n{'='*60}\n"
                f"【文書 No.{doc_num}】\n"
                f"元ファイル: {r.relpath}\n"
                f"{'-'*60}\n"
                f"{r.full_text_for_bind}\n"
                f"{'='*60}\n\n"
            )
            b_len = len(block.encode("utf-8"))
            if current_size + b_len > limit_bytes and current_size > 0:
                flush()
            current_blocks.append(block)
            current_size += b_len
        flush()


def copy_source_files_batched(
    indir: str,
    outdir: str,
    records: List[Record],
    slots_per_batch: int = 46,
) -> Tuple[List[Tuple[str, List[str]]], List[Tuple[str, str]]]:
    """原本PDFを複数のバッチフォルダに分割してコピーする（全件対応）。

    NotebookLMの制限（50MB/件・250MB/ノートブック・50件制限）を守りながら、
    全PDFを収まるだけバッチ分割して漏れなくコピーする。

      - 各バッチフォルダ: 原本コピー_ノートブック01/ 〜
      - 1ファイルあたり 50MB 超はスキップ（サイズ超過として記録）
      - 各バッチ: max slots_per_batch 件 かつ 合計 250MB 以内

    戻り値:
      batches: [(batch_dir, [file_paths]), ...] バッチごとの (フォルダ, ファイルリスト)
      skipped: [(relpath, 理由), ...] スキップしたファイルのリスト
    """
    MAX_FILE_BYTES  = 50  * 1024 * 1024   # 50MB / ファイル
    MAX_BATCH_BYTES = 250 * 1024 * 1024   # 250MB / バッチ

    COPYABLE_EXTS = {".pdf"}

    # 前回の原本コピーフォルダをすべて削除して再生成
    for entry in os.listdir(outdir):
        if entry.startswith("原本コピー") and os.path.isdir(os.path.join(outdir, entry)):
            shutil.rmtree(os.path.join(outdir, entry), ignore_errors=True)

    batches:  List[Tuple[str, List[str]]] = []
    skipped:  List[Tuple[str, str]]       = []
    used_names: set = set()

    current_dir:   Optional[str]   = None
    current_files: List[str]       = []
    current_bytes: int             = 0
    batch_num:     int             = 0

    def _flush():
        nonlocal current_dir, current_files, current_bytes
        if current_dir and current_files:
            batches.append((current_dir, current_files[:]))
        current_dir   = None
        current_files = []
        current_bytes = 0

    def _new_batch() -> str:
        nonlocal batch_num, current_dir, current_files, current_bytes
        _flush()
        batch_num += 1
        d = os.path.join(outdir, f"原本コピー_ノートブック{batch_num:02d}")
        os.makedirs(d, exist_ok=True)
        current_dir   = d
        current_files = []
        current_bytes = 0
        return d

    for r in records:
        if r.ext.lower() not in COPYABLE_EXTS:
            continue
        src = os.path.join(indir, r.relpath)
        if not os.path.isfile(get_safe_path(src)):
            continue

        file_size = os.path.getsize(get_safe_path(src))

        # 50MB 超はどのバッチにも入れられない
        if file_size > MAX_FILE_BYTES:
            skipped.append((r.relpath, f"ファイルサイズ超過 ({file_size // (1024*1024)}MB > 50MB)"))
            continue

        # バッチが未作成、または現バッチが満杯なら新バッチ開始
        need_new = (
            current_dir is None
            or len(current_files) >= slots_per_batch
            or current_bytes + file_size > MAX_BATCH_BYTES
        )
        if need_new:
            _new_batch()

        # フラットなファイル名（パス区切り→アンダースコア）
        safe_name = r.relpath.replace(os.sep, "_").replace("/", "_")
        base, ext = os.path.splitext(safe_name)
        candidate = safe_name
        counter = 1
        while candidate in used_names:
            candidate = f"{base}_{counter}{ext}"
            counter += 1
        used_names.add(candidate)

        dst = os.path.join(current_dir, candidate)
        try:
            shutil.copy2(get_safe_path(src), dst)
            current_files.append(dst)
            current_bytes += file_size
        except Exception:
            pass

    _flush()
    return batches, skipped


def write_notebook_preamble(
    outdir: str,
    records: List[Record],
    bundle_files: List[str],
    copied_files: List[str],
) -> str:
    """NotebookLMへの説明文書を生成する（最初にアップロードするファイル）。

    テキスト抽出ファイルと原本PDFの両方が入っていること、
    矛盾がある場合は原本を優先することを NotebookLM に明示する。
    """
    pdf_records = [r for r in records if r.ext.lower() == ".pdf"]
    ocr_records = [r for r in pdf_records if r.ocr_quality < 1.0]
    type_order = ["法令", "通知", "マニュアル"]
    type_counts: Dict[str, int] = {}
    for r in records:
        type_counts[r.doc_type] = type_counts.get(r.doc_type, 0) + 1

    lines: List[str] = [
        "=" * 60,
        "【このNotebookLMの使い方・注意事項】",
        "（最初にこのファイルを必ずお読みください）",
        "=" * 60,
        "",
        "このノートブックには以下の2種類のソースが格納されています。",
        "",
        "─" * 40,
        "■ テキスト抽出ファイル（NotebookLM用_○○.txt）",
        "─" * 40,
        f"  {len(bundle_files)}ファイル　収録文書 {len(records)}件",
    ]
    for dtype in type_order:
        if dtype in type_counts:
            lines.append(f"  ・{dtype}: {type_counts[dtype]}件")
    lines += [
        "",
        "  複数の文書を1ファイルにまとめたものです。",
        "  PDF・Word・Excel等から自動でテキストを抽出しています。",
    ]
    if ocr_records:
        lines += [
            f"  ※ スキャンPDF {len(ocr_records)}件はOCR読取のため、",
            "    数値・固有名詞等に誤字が含まれる可能性があります。",
        ]
    total_pdf = len(copied_files)
    lines += [
        "",
        "─" * 40,
        "■ 原本PDFファイル（個別にアップロード）",
        "─" * 40,
        f"  合計 {total_pdf}ファイル（複数ノートブックに分割して投入）",
        "",
        "  元のPDFをそのままアップロードしたものです。",
        "  テキスト抽出ファイルと同じ文書の正本です。",
        "",
        "=" * 60,
        "■ 回答時に必ず守ってほしい注意事項",
        "=" * 60,
        "",
        "1. 数値・日付・固有名詞・法令条文番号等は",
        "   必ず原本PDFファイルも参照して確認してください。",
        "",
        "2. テキストファイルと原本PDFで内容が異なる場合は、",
        "   【原本PDFを優先】してください。",
        "   テキストはOCR（機械読取）のため誤字が含まれる場合があります。",
        "",
        "3. どちらのソースを参照したかを回答に必ず明記してください。",
        "   例：「（出典: ○○通知 原本PDF）」",
        "",
        "4. 確認できなかった箇所や不確かな情報には",
        "   「要確認」と付記してください。",
        "",
    ]

    fpath = os.path.join(outdir, "00_はじめに_NotebookLM用.txt")
    with open(fpath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return fpath


def write_upload_guide(
    outdir: str,
    bundle_files: List[str],
    batches: List[Tuple[str, List[str]]],
    skipped_files: Optional[List[Tuple[str, str]]] = None,
):
    """NotebookLMへの投入順序ガイドを生成する（00_投入ガイド.txt）。

    原本PDFは複数バッチに分割済み。各バッチが1つのNotebookLMノートブックに対応する。
    50件制限・50MB/件・250MB合計を守りながら全件を網羅できるよう案内する。
    """
    skipped_files = skipped_files or []
    total_pdf = sum(len(files) for _, files in batches)
    nb_count = len(batches) if batches else 1

    lines: List[str] = [
        "=" * 60,
        "【NotebookLMへの投入ガイド】",
        "=" * 60,
        "",
        f"原本PDF総数: {total_pdf}件  → {nb_count}つのノートブックに分割",
        f"テキストバンドル: {len(bundle_files)}件（全ノートブックに投入）",
        f"50MB超のためスキップ: {len(skipped_files)}件",
        "",
        "━" * 40,
        "【ノートブックの作り方】",
        "━" * 40,
        "",
        "◆ テキスト専用ノートブック（全文検索・要約用）",
        "  ① 00_はじめに_NotebookLM用.txt",
    ]
    for f in bundle_files:
        lines.append(f"  ② {os.path.basename(f)}")
    lines += [""]

    for i, (batch_dir, files) in enumerate(batches, start=1):
        batch_name = os.path.basename(batch_dir)
        batch_mb = sum(
            os.path.getsize(f) for f in files if os.path.isfile(f)
        ) // (1024 * 1024)
        lines += [
            f"◆ ノートブック {i}（原本PDF照合用）← {batch_name}/ フォルダ",
            f"  ファイル数: {len(files)}件  合計サイズ: 約{batch_mb}MB",
            f"  ① 00_はじめに_NotebookLM用.txt（各ノートブックに必ず入れる）",
        ]
        for f in bundle_files:
            lines.append(f"  ② {os.path.basename(f)}（テキストバンドル・全部入れる）")
        lines.append(f"  ③ {batch_name}/ フォルダ内の全PDFをアップロード")
        lines.append("")

    if skipped_files:
        lines += [
            "─" * 40,
            f"【50MB超のためスキップされたPDF: {len(skipped_files)}件】",
            "  ※ NotebookLMのファイルサイズ上限(50MB)を超えているためコピー未実施",
            "─" * 40,
        ]
        for relpath, reason in skipped_files:
            lines.append(f"  スキップ: {os.path.basename(relpath)}  ({reason})")
        lines.append("")

    fpath = os.path.join(outdir, "00_投入ガイド.txt")
    with open(fpath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def write_cross_reference_map(outdir: str, records: List[Record]):
    """相互参照マップを生成する（人間確認用・NotebookLMには入れない）。
    通知が参照する法令条文と、法令文書を紐付ける。
    ★ これは機械推定なので必ず人間が確認すること。"""

    law_records = [r for r in records if r.doc_type == "法令"]
    notice_records = [r for r in records if r.doc_type == "通知"]
    manual_records = [r for r in records if r.doc_type == "マニュアル"]

    lines: List[str] = []
    lines.append("=" * 60)
    lines.append("【相互参照マップ】法令・通知・マニュアルの関連付け")
    lines.append("=" * 60)
    lines.append("")
    lines.append("★★★ 注意 ★★★")
    lines.append("このファイルの内容は機械（テキストパターン）による推定です。")
    lines.append("間違いが含まれている可能性があります。")
    lines.append("必ず人間が確認してから利用してください。")
    lines.append("")
    lines.append("【NotebookLMへの入力について】")
    lines.append("このファイルの内容を確認し、間違いがなければNotebookLMに入れて")
    lines.append("ください。間違いがあればその部分を修正するか、このファイルは")
    lines.append("NotebookLMに入れずに参考資料としてのみ利用してください。")
    lines.append("")
    lines.append(f"法令: {len(law_records)}件 / 通知: {len(notice_records)}件 / マニュアル: {len(manual_records)}件")
    lines.append("")

    # ── 法令一覧 ──
    if law_records:
        lines.append("-" * 40)
        lines.append("■ 収録法令一覧")
        lines.append("-" * 40)
        for r in law_records:
            lines.append(f"  ・{r.title_guess}（{r.date_guess or '日付不明'}）")
        lines.append("")

    # ── 通知→法令の参照関係 ──
    if notice_records:
        lines.append("-" * 40)
        lines.append("■ 通知から法令への参照関係")
        lines.append("-" * 40)
        for r in notice_records:
            if r.related_laws:
                lines.append(f"  [{r.date_guess or '日付不明'}] {r.title_guess}")
                for law_ref in r.related_laws:
                    # 参照先の法令文書が存在するか確認
                    matched_law = ""
                    for lr in law_records:
                        if any(keyword in lr.title_guess for keyword in _extract_law_keywords(law_ref)):
                            matched_law = f" → 収録済み: {lr.title_guess}"
                            break
                    lines.append(f"    → {law_ref}{matched_law}")
                if r.amendments:
                    for a in r.amendments:
                        lines.append(f"    [改廃] {a}")
                lines.append("")

    # ── マニュアル一覧 ──
    if manual_records:
        lines.append("-" * 40)
        lines.append("■ 社内マニュアル一覧")
        lines.append("-" * 40)
        for r in manual_records:
            ref_str = ""
            if r.related_laws:
                ref_str = f" → 関連法令: {', '.join(r.related_laws[:3])}"
            lines.append(f"  ・{r.title_guess}{ref_str}")
        lines.append("")

    # ── 法令条文→通知の逆引き ──
    if law_records and notice_records:
        lines.append("-" * 40)
        lines.append("■ 法令条文から通知への逆引き（どの条文にどの通知が関連するか）")
        lines.append("-" * 40)
        # 法令条文をキーにして通知をグループ化
        law_to_notices: Dict[str, List[str]] = {}
        for r in notice_records:
            for law_ref in r.related_laws:
                if law_ref not in law_to_notices:
                    law_to_notices[law_ref] = []
                law_to_notices[law_ref].append(f"{r.title_guess}（{r.date_guess or ''}）")
        for law_ref, notices in sorted(law_to_notices.items()):
            lines.append(f"  {law_ref}:")
            for notice_title in notices[:10]:
                lines.append(f"    ← {notice_title}")
        lines.append("")

    content = "\n".join(lines)
    with open(os.path.join(outdir, "00_相互参照マップ.txt"), "w", encoding="utf-8") as f:
        f.write(content)


def _extract_law_keywords(law_ref: str) -> List[str]:
    """法令参照文字列からマッチング用キーワードを抽出する"""
    keywords = []
    for name in ["消防法", "政令", "規則", "省令", "条例", "告示"]:
        if name in law_ref:
            keywords.append(name)
    return keywords if keywords else [law_ref[:4]]

def compute_sha1(path: str) -> str:
    """ファイルのSHA1ハッシュを計算して重複ファイル検出に使う"""
    h = hashlib.sha1()
    try:
        with open(get_safe_path(path), "rb") as f:
            for chunk in iter(lambda: f.read(65536), b""):
                h.update(chunk)
        return h.hexdigest()
    except Exception:
        return ""

def extract_txt(path: str) -> Tuple[str, str]:
    """プレーンテキストファイルを読み込む（文字コードを自動判定）"""
    for enc in ("utf-8-sig", "cp932", "utf-8", "latin-1"):
        try:
            with open(get_safe_path(path), "r", encoding=enc, errors="ignore") as f:
                return f.read(), "txt_read"
        except Exception:
            continue
    return "", "txt_err"

def extract_csv(path: str) -> Tuple[str, str]:
    """CSVファイルをMarkdown表形式に整形する"""
    for enc in ("utf-8-sig", "cp932", "utf-8"):
        try:
            with open(get_safe_path(path), "r", encoding=enc, newline="", errors="ignore") as f:
                rows = list(csv.reader(f))
            if not rows:
                return "", "csv_empty"
            out = []
            for row in rows[:400]:
                if any(c.strip() for c in row):
                    out.append("| " + " | ".join([c.strip().replace("\n", " ") for c in row]) + " |")
            return "\n".join(out), "csv_md"
        except Exception:
            continue
    return "", "csv_err"

def write_html_report(outdir: str, records: List[Record]):
    """人間が見やすいHTMLレポートを生成する（ブラウザで開くだけでOK）"""
    def esc(s: object) -> str:
        return _html.escape(str(s) if s is not None else "")

    total       = len(records)
    ok_count    = sum(1 for r in records if not r.needs_review)
    needs_rev_count = total - ok_count
    ok_pct      = round(ok_count    / total * 100) if total else 0
    rev_pct     = round(needs_rev_count / total * 100) if total else 0

    # ─── ファイル種別集計 ─────────────────────────────────────────
    ext_label_map = {
        ".pdf": "PDF", ".docx": "Word",
        ".xlsx": "Excel", ".xlsm": "Excel", ".xls": "Excel",
        ".xdw": "DocuWorks", ".xbd": "DocuWorks",
        ".txt": "テキスト", ".csv": "CSV",
    }
    ext_counts: Dict[str, int] = {}
    for r in records:
        lbl = ext_label_map.get(r.ext.lower(), f"その他({r.ext})")
        ext_counts[lbl] = ext_counts.get(lbl, 0) + 1
    ext_breakdown_parts = [
        f'<span class="type-chip">{esc(lbl)} <b>{cnt}</b>件</span>'
        for lbl, cnt in sorted(ext_counts.items(), key=lambda x: -x[1])
    ]
    ext_breakdown_html = "".join(ext_breakdown_parts)

    # ─── 抽出方式集計（抽出方式別テーブル） ─────────────────────────
    method_counts: Dict[str, int] = {}
    for r in records:
        method_counts[r.method] = method_counts.get(r.method, 0) + 1
    method_rows = "".join(
        f"<tr><td>{esc(m)}</td><td class='mcnt'>{c}</td></tr>"
        for m, c in sorted(method_counts.items(), key=lambda x: -x[1])
    )

    # ─── 要確認の主要理由を集計 ─────────────────────────────────────
    review_reasons: Dict[str, int] = {}
    for r in records:
        if r.needs_review and r.reason:
            key = r.reason[:35] + ("…" if len(r.reason) > 35 else "")
            review_reasons[key] = review_reasons.get(key, 0) + 1
    review_reason_rows = "".join(
        f'<li><span class="rr-count">{c}件</span> {esc(k)}</li>'
        for k, c in sorted(review_reasons.items(), key=lambda x: -x[1])[:5]
    )

    # ─── 文書タイプ別集計 ────────────────────────────────────────────
    dtype_counts: Dict[str, int] = {}
    for r in records:
        dtype_counts[r.doc_type] = dtype_counts.get(r.doc_type, 0) + 1
    _dtype_css = {"法令": "law", "通知": "notice", "マニュアル": "manual"}
    dtype_breakdown_parts = [
        f'<span class="type-chip dtype-{_dtype_css.get(dt, "notice")}">{esc(dt)} <b>{cnt}</b>件</span>'
        for dt, cnt in [("法令", dtype_counts.get("法令", 0)),
                        ("通知", dtype_counts.get("通知", 0)),
                        ("マニュアル", dtype_counts.get("マニュアル", 0))]
        if cnt > 0
    ]
    dtype_breakdown_html = "".join(dtype_breakdown_parts)

    # ─── バッジ色 ─────────────────────────────────────────────────
    FAC_COLOR  = "#2563eb"
    WORK_COLOR = "#16a34a"
    def make_badge(text: str, color: str) -> str:
        return f'<span class="badge" style="background:{color}">{esc(text)}</span>'

    # ─── TOCアイテム生成 ─────────────────────────────────────────
    toc_items_html: List[str] = []
    for idx, r in enumerate(records):
        toc_cls  = "toc-review" if r.needs_review else "toc-ok"
        toc_icon = "⚠" if r.needs_review else "✓"
        short_t  = r.title_guess[:40] + ("…" if len(r.title_guess) > 40 else "")
        d_str    = r.date_guess or "日付不明"
        tsearch  = (r.title_guess + " " + d_str).lower().replace('"', "")
        toc_items_html.append(
            f'<a href="#card-{idx}" class="toc-item {toc_cls}" data-search="{esc(tsearch)}">'
            f'<span class="toc-icon">{toc_icon}</span>'
            f'<span class="toc-body">'
            f'<span class="toc-num">{idx + 1}.</span>'
            f'<span class="toc-title">{esc(short_t)}</span>'
            f'<span class="toc-date">{esc(d_str)}</span>'
            f'</span></a>'
        )

    # ─── カード生成 ───────────────────────────────────────────────
    cards_html: List[str] = []
    for idx, r in enumerate(records):
        card_cls  = "card-review" if r.needs_review else "card-ok"
        rev_badge = '<span class="rev-badge">⚠ 要確認</span>' if r.needs_review else \
                    '<span class="ok-badge">✓ 正常</span>'
        fac_badges  = "".join(make_badge(t, FAC_COLOR)  for t in r.tags_facility)
        work_badges = "".join(make_badge(t, WORK_COLOR) for t in r.tags_work)
        tags_html   = (fac_badges + work_badges) or \
                      '<span style="color:#94a3b8;font-size:12px">タグなし</span>'
        date_str   = esc(r.date_guess)   or "日付不明"
        issuer_str = esc(r.issuer_guess) or "発出者不明"
        pages_str  = f"/{r.pages}p" if r.pages else ""
        size_kb    = f"{r.size // 1024:,} KB" if r.size >= 1024 else f"{r.size} B"
        reason_html = (
            f'<div class="reason-box">⚠ {esc(r.reason)}</div>' if r.reason else ""
        )

        # 文書タイプバッジ
        dtype_cls = {"法令": "dtype-law", "通知": "dtype-notice", "マニュアル": "dtype-manual"}.get(r.doc_type, "dtype-notice")
        dtype_badge_html = f'<span class="dtype-badge {dtype_cls}">{esc(r.doc_type)}</span>'

        # OCR品質バッジ（OCR処理したファイルのみ表示）
        ocr_badge_html = ""
        if r.ocr_quality < 1.0:
            if r.ocr_quality >= 0.6:
                ocr_badge_html = f'<span class="ocr-badge ocr-ok">OCR品質: {r.ocr_quality:.0%}</span>'
            elif r.ocr_quality >= 0.35:
                ocr_badge_html = f'<span class="ocr-badge ocr-warn">OCR品質: {r.ocr_quality:.0%}</span>'
            else:
                ocr_badge_html = f'<span class="ocr-badge ocr-bad">OCR品質: {r.ocr_quality:.0%}</span>'

        # 改廃情報（検出された場合のみ）
        amend_html = ""
        if r.amendments:
            amend_items = "".join(f'<span class="amend-chip">{esc(a)}</span>' for a in r.amendments[:3])
            amend_html = f'<div class="amend-row">改廃: {amend_items}</div>'

        # 関連法令（検出された場合のみ）
        laws_html = ""
        if r.related_laws:
            law_items = "".join(f'<span class="law-chip">{esc(l)}</span>' for l in r.related_laws[:5])
            laws_html = f'<div class="law-row">関連法令: {law_items}</div>'

        search_data = " ".join([
            r.title_guess, r.summary, r.relpath,
            r.date_guess, r.issuer_guess, r.doc_type,
            " ".join(r.tags_facility), " ".join(r.tags_work),
            " ".join(r.related_laws), " ".join(r.amendments),
            r.reason, r.method,
        ]).replace('"', '')
        summary_html = (esc(r.summary)
                        or '<i style="color:#94a3b8">本文を抽出できませんでした</i>')
        cards_html.append(f"""
<div id="card-{idx}" class="card {card_cls}" data-search="{esc(search_data.lower())}">
  <div class="card-header">
    <div class="card-title">{esc(r.title_guess)}</div>
    <div class="card-badges">{dtype_badge_html}{ocr_badge_html}{rev_badge}</div>
  </div>
  <div class="meta">
    <span>📅 {date_str}</span>
    <span>🏢 {issuer_str}</span>
    <span>📄 {esc(r.ext.upper().lstrip('.'))}{pages_str} · {size_kb}</span>
    <span class="method-tag">抽出: {esc(r.method)}</span>
  </div>
  <div class="tags">{tags_html}</div>
  {amend_html}
  {laws_html}
  <div class="summary">{summary_html}</div>
  <div class="filepath">📁 {esc(r.relpath)}</div>
  {reason_html}
</div>""")

    gen_time = time.strftime('%Y年%m月%d日 %H:%M:%S')

    html_content = f"""<!DOCTYPE html>
<html lang="ja">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>NoticeForge 処理レポート</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:'Meiryo UI','Yu Gothic UI','Hiragino Sans',sans-serif;background:#f1f5f9;color:#1e293b;font-size:14px}}

/* ════════════════════════════════════
   左サイドバー（文書目次）
   ════════════════════════════════════ */
.toc-sidebar{{
  position:fixed;left:0;top:0;width:300px;height:100vh;
  background:#0f172a;color:#e2e8f0;
  display:flex;flex-direction:column;z-index:200;
  border-right:1px solid #1e3a5f;
}}
.toc-head{{
  padding:14px 16px;font-size:14px;font-weight:bold;
  background:#1e3a8a;color:white;
  display:flex;align-items:center;gap:8px;flex-shrink:0;
}}
.toc-summary-row{{
  padding:8px 16px;font-size:12px;color:#94a3b8;
  background:#1e293b;border-bottom:1px solid #334155;flex-shrink:0;
  display:flex;gap:14px;
}}
.toc-ok-sum{{color:#4ade80;font-weight:bold}}
.toc-rev-sum{{color:#f87171;font-weight:bold}}
.toc-filter-wrap{{
  padding:8px 12px;background:#1e293b;
  border-bottom:1px solid #334155;flex-shrink:0;
}}
.toc-filter{{
  width:100%;padding:6px 10px;border-radius:6px;
  border:1px solid #334155;background:#0f172a;
  color:#e2e8f0;font-size:12px;font-family:inherit;outline:none;
}}
.toc-filter:focus{{border-color:#3b82f6}}
.toc-nav{{flex:1;overflow-y:auto;padding:4px 0}}
.toc-nav::-webkit-scrollbar{{width:4px}}
.toc-nav::-webkit-scrollbar-thumb{{background:#334155;border-radius:2px}}
.toc-item{{
  display:flex;align-items:flex-start;gap:8px;
  padding:7px 14px;text-decoration:none;color:#cbd5e1;
  font-size:12px;line-height:1.4;
  border-left:3px solid transparent;
  transition:background .15s,border-color .15s;
}}
.toc-item:hover{{background:#1e293b;color:white}}
.toc-item.active{{background:#1e3a8a;border-left-color:#60a5fa;color:white}}
.toc-icon{{font-size:11px;flex-shrink:0;margin-top:1px;width:14px;text-align:center}}
.toc-ok   .toc-icon{{color:#4ade80}}
.toc-review .toc-icon{{color:#f87171}}
.toc-body{{display:flex;flex-direction:column;min-width:0;flex:1}}
.toc-num{{color:#64748b;font-size:10px}}
.toc-title{{font-size:12px;color:inherit;white-space:normal;overflow:hidden;display:-webkit-box;-webkit-line-clamp:2;-webkit-box-orient:vertical}}
.toc-date{{font-size:10px;color:#64748b;margin-top:1px}}
.toc-item.toc-hidden{{display:none}}
.toc-empty{{padding:16px;font-size:12px;color:#475569;text-align:center}}

/* ════════════════════════════════════
   メインコンテンツ
   ════════════════════════════════════ */
.main-wrapper{{margin-left:300px}}

/* ─── ページヘッダー ─── */
.page-header{{
  background:linear-gradient(135deg,#1e40af,#2563eb);
  color:white;padding:20px 32px;
  display:flex;justify-content:space-between;align-items:flex-end;
  flex-wrap:wrap;gap:8px;
}}
.page-header h1{{font-size:22px;font-weight:bold}}
.page-header .sub{{opacity:.75;font-size:12px;margin-top:4px}}

/* ─── 処理概要セクション ─── */
.overview-section{{
  background:white;border-bottom:1px solid #e2e8f0;padding:20px 32px 16px;
}}
.overview-title{{
  font-size:13px;font-weight:bold;color:#64748b;
  text-transform:uppercase;letter-spacing:.05em;margin-bottom:14px;
}}
.stats-row{{display:flex;gap:12px;flex-wrap:wrap;margin-bottom:16px;align-items:stretch}}
.stat-box{{
  background:#f8fafc;border:1px solid #e2e8f0;border-radius:10px;
  padding:14px 24px;text-align:center;min-width:110px;
}}
.stat-box .num{{font-size:30px;font-weight:bold;color:#1e40af;line-height:1}}
.stat-box .lbl{{font-size:11px;color:#64748b;margin-top:6px}}
.stat-box .pct{{font-size:11px;color:#94a3b8;margin-top:2px}}
.stat-box.warn .num{{color:#dc2626}}
.stat-box.good .num{{color:#16a34a}}
.overview-bottom{{display:flex;gap:24px;flex-wrap:wrap;align-items:flex-start}}
.type-section{{flex:1;min-width:200px}}
.type-label{{font-size:12px;color:#64748b;font-weight:bold;margin-bottom:8px}}
.type-chips{{display:flex;gap:8px;flex-wrap:wrap}}
.type-chip{{
  background:#f1f5f9;border:1px solid #e2e8f0;border-radius:20px;
  padding:4px 12px;font-size:12px;color:#475569;
}}
.type-chip b{{color:#1e40af}}
.method-section{{flex:1;min-width:180px}}
.method-section table{{font-size:12px;border-collapse:collapse;width:100%}}
.method-section td{{padding:3px 8px;border-bottom:1px solid #f1f5f9;color:#475569}}
.method-section td.mcnt{{text-align:right;font-weight:bold;color:#1e40af}}
.method-section tr:last-child td{{border-bottom:none}}
.review-section{{flex:1;min-width:180px}}
.review-reasons{{list-style:none;font-size:12px;color:#92400e}}
.review-reasons li{{padding:2px 0;display:flex;align-items:baseline;gap:6px}}
.rr-count{{
  background:#fee2e2;color:#dc2626;border-radius:4px;
  padding:1px 6px;font-weight:bold;font-size:11px;white-space:nowrap;flex-shrink:0;
}}
.guide-box{{
  background:#eff6ff;border:1px solid #bfdbfe;border-radius:8px;
  padding:10px 16px;font-size:12px;color:#1e40af;margin-top:14px;
  display:flex;align-items:flex-start;gap:8px;
}}
.guide-box strong{{font-weight:bold}}

/* ─── 検索バー（sticky）─── */
.search-bar{{
  background:white;padding:10px 24px;border-bottom:1px solid #e2e8f0;
  display:flex;align-items:center;gap:10px;
  position:sticky;top:0;z-index:100;
  box-shadow:0 2px 6px rgba(0,0,0,.06);
}}
.search-input{{
  flex:1;max-width:680px;padding:9px 14px 9px 40px;
  border:2px solid #e2e8f0;border-radius:8px;
  font-size:13px;font-family:inherit;outline:none;
  transition:border-color .2s;
  background:url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='16' height='16' fill='none' stroke='%2394a3b8' stroke-width='2' viewBox='0 0 24 24'%3E%3Ccircle cx='11' cy='11' r='8'/%3E%3Cpath d='m21 21-4.35-4.35'/%3E%3C/svg%3E") no-repeat 12px center;
}}
.search-input:focus{{border-color:#2563eb}}
.search-hint{{font-size:11px;color:#94a3b8}}
.search-count{{font-size:13px;color:#64748b;font-weight:bold;white-space:nowrap;margin-left:auto}}
.no-results{{text-align:center;padding:64px 16px;color:#94a3b8;font-size:15px;display:none}}

/* ─── カード ─── */
.container{{max-width:1000px;margin:20px auto;padding:0 20px}}
.card{{
  background:white;border-radius:10px;padding:18px 22px;margin-bottom:14px;
  border-left:5px solid #94a3b8;
  box-shadow:0 1px 4px rgba(0,0,0,.07);
  transition:box-shadow .2s;scroll-margin-top:56px;
}}
.card:hover{{box-shadow:0 3px 10px rgba(0,0,0,.12)}}
.card.highlight{{outline:3px solid #3b82f6;outline-offset:2px}}
.card-ok{{border-left-color:#16a34a}}
.card-review{{border-left-color:#dc2626}}
.card-header{{display:flex;justify-content:space-between;align-items:flex-start;gap:12px;margin-bottom:10px}}
.card-title{{font-size:15px;font-weight:bold;color:#0f172a;line-height:1.5;flex:1}}
.ok-badge{{background:#dcfce7;color:#16a34a;border:1px solid #86efac;border-radius:6px;padding:2px 10px;font-size:12px;font-weight:bold;white-space:nowrap}}
.rev-badge{{background:#fee2e2;color:#dc2626;border:1px solid #fca5a5;border-radius:6px;padding:2px 10px;font-size:12px;font-weight:bold;white-space:nowrap}}
.meta{{display:flex;gap:14px;flex-wrap:wrap;color:#64748b;font-size:12px;margin-bottom:10px}}
.method-tag{{color:#94a3b8;font-size:11px}}
.tags{{display:flex;gap:6px;flex-wrap:wrap;margin-bottom:12px}}
.badge{{color:white;padding:2px 10px;border-radius:12px;font-size:12px;font-weight:500}}
.summary{{
  background:#f8fafc;border:1px solid #e2e8f0;border-radius:6px;
  padding:10px 14px;font-size:13px;line-height:1.8;color:#334155;
  max-height:160px;overflow-y:auto;margin-bottom:10px;white-space:pre-wrap;
}}
.filepath{{font-size:11px;color:#94a3b8;font-family:'Consolas','Courier New',monospace;word-break:break-all}}
.reason-box{{margin-top:8px;font-size:12px;color:#92400e;background:#fffbeb;border:1px solid #fde68a;border-radius:5px;padding:6px 12px}}
.card-badges{{display:flex;gap:6px;align-items:center;flex-shrink:0}}
.ocr-badge{{border-radius:6px;padding:2px 8px;font-size:11px;font-weight:bold;white-space:nowrap}}
.ocr-ok{{background:#dcfce7;color:#16a34a;border:1px solid #86efac}}
.ocr-warn{{background:#fef3c7;color:#d97706;border:1px solid #fcd34d}}
.ocr-bad{{background:#fee2e2;color:#dc2626;border:1px solid #fca5a5}}
.amend-row,.law-row{{font-size:12px;color:#475569;margin-bottom:8px;display:flex;gap:6px;flex-wrap:wrap;align-items:center}}
.amend-chip{{background:#fef3c7;color:#92400e;border:1px solid #fde68a;border-radius:4px;padding:1px 8px;font-size:11px}}
.law-chip{{background:#ede9fe;color:#6d28d9;border:1px solid #c4b5fd;border-radius:4px;padding:1px 8px;font-size:11px}}
.dtype-badge{{border-radius:6px;padding:2px 10px;font-size:11px;font-weight:bold;white-space:nowrap}}
.dtype-law{{background:#dbeafe;color:#1d4ed8;border:1px solid #93c5fd}}
.dtype-notice{{background:#f0fdf4;color:#15803d;border:1px solid #86efac}}
.dtype-manual{{background:#fef3c7;color:#92400e;border:1px solid #fcd34d}}

/* ─── フッター ─── */
.footer{{text-align:center;color:#94a3b8;font-size:11px;padding:24px;margin-top:8px}}

/* ─── レスポンシブ（狭い画面では目次非表示） ─── */
@media(max-width:900px){{
  .toc-sidebar{{display:none}}
  .main-wrapper{{margin-left:0}}
}}
</style>
</head>
<body>

<!-- ════ 左サイドバー（文書目次）════ -->
<aside class="toc-sidebar">
  <div class="toc-head">📋 文書目次</div>
  <div class="toc-summary-row">
    <span class="toc-ok-sum">✓ 正常 {ok_count}件</span>
    <span class="toc-rev-sum">⚠ 要確認 {needs_rev_count}件</span>
  </div>
  <div class="toc-filter-wrap">
    <input class="toc-filter" id="tocFilter" type="text"
      placeholder="目次を絞り込む…" oninput="filterToc()">
  </div>
  <nav class="toc-nav" id="tocNav">
    {''.join(toc_items_html)}
    <div class="toc-empty" id="tocEmpty" style="display:none">該当なし</div>
  </nav>
</aside>

<!-- ════ メインコンテンツ ════ -->
<div class="main-wrapper">

  <!-- ページヘッダー -->
  <header class="page-header">
    <div>
      <h1>NoticeForge 処理レポート</h1>
      <div class="sub">生成日時: {gen_time}</div>
    </div>
  </header>

  <!-- 処理概要 -->
  <section class="overview-section">
    <div class="overview-title">処理概要</div>
    <div class="stats-row">
      <div class="stat-box">
        <div class="num">{total}</div>
        <div class="lbl">総ファイル数</div>
      </div>
      <div class="stat-box good">
        <div class="num">{ok_count}</div>
        <div class="lbl">正常抽出</div>
        <div class="pct">{ok_pct}%</div>
      </div>
      <div class="stat-box warn">
        <div class="num">{needs_rev_count}</div>
        <div class="lbl">要確認</div>
        <div class="pct">{rev_pct}%</div>
      </div>
    </div>
    <div class="overview-bottom">
      <div class="type-section">
        <div class="type-label">文書タイプ</div>
        <div class="type-chips">{dtype_breakdown_html}</div>
      </div>
      <div class="type-section">
        <div class="type-label">ファイル種別</div>
        <div class="type-chips">{ext_breakdown_html}</div>
      </div>
      <div class="method-section">
        <div class="type-label">抽出方式別</div>
        <table><tbody>{method_rows}</tbody></table>
      </div>
      {'<div class="review-section"><div class="type-label">要確認の主な理由</div><ul class="review-reasons">' + review_reason_rows + '</ul></div>' if review_reason_rows else ''}
    </div>
    <div class="guide-box">
      💡 <span><strong>NotebookLMへの入力：</strong>
      「NotebookLM用_○○.txt」を全てアップロードしてください（これが本文です）。
      「00_相互参照マップ.txt」は<strong>機械推定</strong>なので、
      中身を確認して間違いがなければ入れてください。間違いがあれば入れないでください。
      NotebookLMはソースの内容をそのまま引用するため、間違った情報を入れると誤った回答の原因になります。</span>
    </div>
  </section>

  <!-- 検索バー（sticky）-->
  <div class="search-bar">
    <input class="search-input" id="searchInput" type="text"
      placeholder="キーワードで絞り込む（タイトル・発出者・ファイル名・概要など。NotebookLMの引用文をそのまま貼り付けてもOK）"
      oninput="filterCards()">
    <span class="search-hint">→ 元ファイルを素早く特定できます</span>
    <span class="search-count" id="searchCount"></span>
  </div>

  <!-- カード一覧 -->
  <div class="container">
    {''.join(cards_html)}
    <div class="no-results" id="noResults">
      該当するファイルが見つかりませんでした。別のキーワードを試してください。
    </div>
  </div>

  <div class="footer">NoticeForge &mdash; NotebookLM 連携ツール &nbsp;|&nbsp; 生成: {gen_time}</div>
</div>

<script>
/* ── カード検索 ── */
function filterCards() {{
  var q = document.getElementById('searchInput').value.toLowerCase();
  var cards = document.querySelectorAll('.card');
  var shown = 0;
  cards.forEach(function(card) {{
    var match = !q || card.getAttribute('data-search').includes(q);
    card.style.display = match ? '' : 'none';
    if (match) shown++;
  }});
  var countEl = document.getElementById('searchCount');
  var noRes   = document.getElementById('noResults');
  countEl.textContent = q ? (shown + ' 件 / ' + cards.length + ' 件中') : (cards.length + ' 件');
  noRes.style.display  = (q && shown === 0) ? 'block' : 'none';
}}

/* ── 目次絞り込み ── */
function filterToc() {{
  var q = document.getElementById('tocFilter').value.toLowerCase();
  var items = document.querySelectorAll('.toc-item');
  var shown = 0;
  items.forEach(function(a) {{
    var match = !q || a.getAttribute('data-search').includes(q);
    a.classList.toggle('toc-hidden', !match);
    if (match) shown++;
  }});
  document.getElementById('tocEmpty').style.display = (q && shown === 0) ? 'block' : 'none';
}}

/* ── スクロール連動でTOCをハイライト ── */
(function() {{
  var tocItems = {{}};
  document.querySelectorAll('.toc-item').forEach(function(a) {{
    var id = a.getAttribute('href').slice(1);
    tocItems[id] = a;
  }});
  var observer = new IntersectionObserver(function(entries) {{
    entries.forEach(function(entry) {{
      if (entry.isIntersecting) {{
        Object.values(tocItems).forEach(function(a) {{ a.classList.remove('active'); }});
        var active = tocItems[entry.target.id];
        if (active) {{
          active.classList.add('active');
          var nav = document.getElementById('tocNav');
          if (nav) {{
            var offset = active.offsetTop - nav.offsetTop;
            nav.scrollTop = offset - nav.clientHeight / 3;
          }}
        }}
      }}
    }});
  }}, {{ rootMargin: '-5% 0% -70% 0%', threshold: 0 }});
  document.querySelectorAll('.card').forEach(function(c) {{ observer.observe(c); }});

  /* ── 初期件数表示 ── */
  document.getElementById('searchCount').textContent =
    document.querySelectorAll('.card').length + ' 件';

  /* ── TOCリンクをクリックしたときカードを一瞬ハイライト ── */
  document.querySelectorAll('.toc-item').forEach(function(a) {{
    a.addEventListener('click', function() {{
      var id = a.getAttribute('href').slice(1);
      var card = document.getElementById(id);
      if (card) {{
        card.classList.add('highlight');
        setTimeout(function() {{ card.classList.remove('highlight'); }}, 1200);
      }}
    }});
  }});
}})();
</script>
</body>
</html>"""

    with open(os.path.join(outdir, "00_人間用レポート.html"), "w", encoding="utf-8") as f:
        f.write(html_content)


def process_folder(indir: str, outdir: str, cfg: Dict[str, object], progress_callback: Optional[Callable[[int, int, str, str], None]] = None, stop_event=None) -> Tuple[int, int, str]:
    os.makedirs(outdir, exist_ok=True)
    outdir_abs = os.path.abspath(outdir)

    # 前回の生成ファイルを削除（古いデータがNotebookLMに混入しないように）
    # ※ 00_manifest.json だけは差分処理のために残す
    for fname in os.listdir(outdir):
        if fname.startswith("NotebookLM用_") and fname.endswith(".txt"):
            try: os.remove(os.path.join(outdir, fname))
            except Exception: pass
    for fname in (
        "00_統合目次.md", "00_統合目次.xlsx", "00_人間用レポート.html",
        "00_処理ログ.txt", "00_相互参照マップ.txt",
        "00_はじめに_NotebookLM用.txt", "00_投入ガイド.txt",
    ):
        p = os.path.join(outdir, fname)
        if os.path.exists(p):
            try: os.remove(p)
            except Exception: pass
    # 原本コピーフォルダも再生成する（前回分を削除）
    for _entry in os.listdir(outdir):
        if _entry.startswith("原本コピー") and os.path.isdir(os.path.join(outdir, _entry)):
            shutil.rmtree(os.path.join(outdir, _entry), ignore_errors=True)

    max_depth = int(cfg.get("max_depth", 30))
    split_kws = list(cfg.get("main_attach_split_keywords", []))
    min_chars = int(cfg.get("min_chars_mainbody", 400))
    use_ocr = bool(cfg.get("use_ocr", False))
    limit_bytes = int(cfg.get("bind_bytes_limit", 15000000))

    SKIP_FILENAMES = frozenset({"thumbs.db", "desktop.ini", ".ds_store"})
    SKIP_EXTENSIONS = frozenset({".db", ".tmp", ".bak", ".lnk", ".ini", ".cache"})

    # 【バグ修正】出力フォルダが入力フォルダ内にある場合、スキャン対象から除外する
    targets: List[str] = []
    for root, dirs, files in os.walk(indir):
        # 出力フォルダのサブツリーを丸ごとスキップ（dirs を in-place 変更）
        dirs[:] = [
            d for d in dirs
            if os.path.abspath(os.path.join(root, d)) != outdir_abs
        ]
        # 深さ制限
        rel_root = os.path.relpath(root, indir)
        depth = 0 if rel_root == "." else rel_root.count(os.sep) + 1
        if depth >= max_depth:
            dirs.clear()
            continue
        for fn in files:
            if fn.lower() in SKIP_FILENAMES: continue
            if os.path.splitext(fn)[1].lower() in SKIP_EXTENSIONS: continue
            if fn.startswith("~$"): continue
            targets.append(os.path.join(root, fn))

    total_files = len(targets)
    records: List[Record] = []
    seen_sha1: set = set()
    skipped_dup = 0
    skipped_cache = 0

    # マニフェスト（処理キャッシュ）を読み込む
    # → 変更のないファイルは再処理をスキップし、前回結果を再利用する
    # ※ キャッシュバージョンが不一致の場合は全件再処理（概要ロジック変更時の整合性保証）
    manifest_path = os.path.join(outdir, "00_manifest.json")
    manifest: Dict[str, dict] = {}
    if os.path.exists(manifest_path):
        try:
            with open(manifest_path, "r", encoding="utf-8") as f:
                manifest_raw = json.load(f)
            # キャッシュバージョンチェック
            if manifest_raw.get("_cache_version") == _CACHE_VERSION:
                manifest = {k: v for k, v in manifest_raw.items() if k != "_cache_version"}
            else:
                manifest = {}  # バージョン不一致 → 全件再処理
        except Exception:
            manifest = {}

    log_lines: List[str] = [
        "=== NoticeForge 処理ログ ===",
        f"処理日時: {time.strftime('%Y年%m月%d日 %H:%M:%S')}",
        f"入力フォルダ: {indir}",
        f"出力フォルダ: {outdir}",
        f"キャッシュ読込: {len(manifest)} 件",
        "",
        "--- 各ファイルの処理結果 ---",
    ]

    for i, path in enumerate(targets):
        # 停止リクエストをチェック
        if stop_event and stop_event.is_set():
            log_lines.append("[STOPPED] ユーザーにより処理を途中で停止しました。")
            break

        rel = os.path.relpath(path, indir)
        ext = os.path.splitext(path)[1].lower()
        if progress_callback: progress_callback(i + 1, total_files, rel, "(確認中...)")

        sha1 = compute_sha1(path)

        # 重複ファイルチェック
        if sha1 and sha1 in seen_sha1:
            if progress_callback: progress_callback(i + 1, total_files, rel, "(重複・スキップ)")
            log_lines.append(f"[重複スキップ] {rel}")
            skipped_dup += 1
            continue

        # キャッシュヒットチェック（SHA1が一致 → 内容変更なし → 前回結果を再利用）
        if sha1 and sha1 in manifest:
            try:
                cached = manifest[sha1]
                record = Record(**{**cached, "relpath": rel, "sha1": sha1})
                records.append(record)
                seen_sha1.add(sha1)
                if progress_callback: progress_callback(i + 1, total_files, rel, "(キャッシュ使用)")
                log_lines.append(f"[キャッシュ] {rel}")
                skipped_cache += 1
                continue
            except Exception:
                pass  # キャッシュが壊れていたら通常処理にフォールバック

        seen_sha1.add(sha1)
        if progress_callback: progress_callback(i + 1, total_files, rel, "(抽出中...)")

        text, method, reason, pages = "", "unhandled", "", None

        try:
            if ext == ".pdf":
                if use_ocr and progress_callback: progress_callback(i + 1, total_files, rel, "(OCR処理中...時間がかかります)")
                text, pages, method = extract_pdf(path, use_ocr)
            elif ext == ".docx":
                text, method = extract_docx(path)
            elif ext in (".xlsx", ".xlsm", ".xls"):
                text, method = extract_excel(path)
            elif ext in (".xdw", ".xbd"):
                text, method = extract_xdw(path)
            elif ext == ".txt":
                text, method = extract_txt(path)
            elif ext == ".csv":
                text, method = extract_csv(path)
        except Exception as e:
            method, reason = "error", f"抽出エラー: {e.__class__.__name__}"

        text = convert_japanese_year(text)
        main, attach = split_main_attach(text, split_kws)

        # ── 文書タイプ自動判別 ──
        doc_type = _detect_doc_type(rel, main or text)

        # OCR品質スコアを計算（OCR系メソッドのみ）
        ocr_q = 1.0
        if "ocr" in method:
            ocr_q = _compute_ocr_quality(text)

        # 日付のみ抽出（ソート用）
        date_guess = guess_date(text)
        date_sort = _date_to_sort_key(date_guess)

        # ファイルサイズを取得（needs_review判定で使用）
        file_size = os.path.getsize(get_safe_path(path))
        text_len = len(main or text)

        needs_rev = False
        if method in ("unhandled", "error") or "missing" in method:
            needs_rev = True
            if not reason:
                if "xdw_text_extractor_missing" in method:
                    if XDWLIB_AVAILABLE:
                        reason = "DocuWorks Viewer Light は検出済みですが、このファイルのテキスト抽出に失敗しました（文書が保護されている可能性）"
                    else:
                        reason = "DocuWorks テキスト抽出ツールが見つかりません。DocuWorks Viewer Light がインストール済みの場合は xdoc2txt.exe を追加してください: https://ebstudio.info/home/xdoc2txt.html"
                elif method == "unhandled":
                    reason = f"未対応ファイル形式 ({ext})"
                elif "pymupdf_missing" in method:
                    reason = "PyMuPDFが未インストール（pip install PyMuPDF）"
                elif "excel_lib_missing" in method:
                    reason = "Excelライブラリが未インストール（pip install openpyxl xlrd）"
                else:
                    reason = f"抽出失敗: {method}"
        elif ext in (".xlsx", ".xlsm", ".xls", ".csv", ".txt"):
            pass
        elif text_len < 30:
            needs_rev = True
            if ext == ".pdf" and not TESSERACT_AVAILABLE:
                reason = "画像PDFの可能性（Tesseract OCRが未インストールのため読取不可）"
            elif ext == ".pdf":
                reason = "OCRを試みましたが読取できませんでした（スキャン品質が低い可能性）"
            else:
                reason = f"本文がほぼ空です（{text_len}文字）"
        elif file_size > 30000 and text_len < 100:
            needs_rev = True
            reason = f"ファイルサイズ({file_size // 1024}KB)に対して本文が短すぎます（{text_len}文字・画像PDF等の可能性）"

        # OCR品質が低い場合も要確認
        if ocr_q < 0.35 and not needs_rev:
            needs_rev = True
            reason = f"OCR品質が低い（スコア: {ocr_q}）。元ファイルの目視確認を推奨"

        # ── ペイロード（NotebookLM用テキスト）──
        # ★重要: NotebookLMに渡すテキストにはAI推定情報を入れない
        # NotebookLMは入力ソースだけを参照するため、推定が間違っていると
        # NotebookLMが誤情報を「事実」として引用してしまう。
        # タイトル・日付・発出者は本文中に元々含まれているのでそのまま渡す。
        payload = f"# 本文\n{main.strip()}"
        if attach.strip():
            payload += f"\n\n# 添付資料\n{attach.strip()}"

        log_lines.append(f"[{method}][{doc_type}] {rel}" + (f"  OCR品質:{ocr_q}" if ocr_q < 1.0 else ""))
        if reason:
            log_lines.append(f"  → {reason}")

        records.append(Record(
            relpath=rel, ext=ext,
            size=file_size,
            mtime=os.path.getmtime(get_safe_path(path)),
            sha1=sha1, method=method, pages=pages,
            text_chars=len(text), needs_review=needs_rev, reason=reason,
            title_guess="", date_guess=date_guess, issuer_guess="",
            summary="", tags_facility=[], tags_work=[], tag_evidence={},
            out_txt="", full_text_for_bind=payload,
            doc_type=doc_type,
            ocr_quality=ocr_q, related_laws=[], amendments=[],
            date_sort_key=date_sort,
        ))

    # ── タイプ別＋時系列ソート（法令→通知→マニュアル、各タイプ内は日付新しい順）──
    type_sort_order = {"法令": 0, "通知": 1, "マニュアル": 2}
    records.sort(key=lambda r: (type_sort_order.get(r.doc_type, 9), r.date_sort_key), reverse=False)
    # 日付は新しい順にしたいので、タイプ内で逆順にする
    records.sort(key=lambda r: type_sort_order.get(r.doc_type, 9))
    # タイプ別にグループ化してから日付ソート
    sorted_records: List[Record] = []
    for dtype in ["法令", "通知", "マニュアル"]:
        group = [r for r in records if r.doc_type == dtype]
        group.sort(key=lambda r: r.date_sort_key, reverse=True)
        sorted_records.extend(group)
    records[:] = sorted_records

    write_binded_texts(outdir, records, limit_bytes)

    # 原本PDFをコピーし、説明文書・投入ガイドを生成する
    import glob as _glob
    bundle_files = sorted(_glob.glob(os.path.join(outdir, "NotebookLM用_*.txt")))
    # 各バッチで使えるスロット数（説明文書1件 + バンドル分を除いた残り）
    _pdf_slots = max(50 - 1 - len(bundle_files), 0)
    batches, skipped_files = copy_source_files_batched(indir, outdir, records, slots_per_batch=_pdf_slots)
    all_copied = [f for _, files in batches for f in files]
    write_notebook_preamble(outdir, records, bundle_files, all_copied)
    write_upload_guide(outdir, bundle_files, batches, skipped_files)

    # サマリーを集計してログファイルに保存
    needs_rev_count = len([r for r in records if r.needs_review])
    review_breakdown: Dict[str, int] = {}
    for r in records:
        if r.needs_review:
            # 理由の先頭部分（40文字まで）をキーにして集計
            key = r.reason[:40] if r.reason else r.method
            review_breakdown[key] = review_breakdown.get(key, 0) + 1

    # 文書タイプ別集計
    dtype_log: Dict[str, int] = {}
    for r in records:
        dtype_log[r.doc_type] = dtype_log.get(r.doc_type, 0) + 1

    log_lines += [
        "",
        "--- サマリー ---",
        f"総処理数: {len(records)} 件（うちキャッシュ利用: {skipped_cache} 件）",
        f"文書タイプ別: " + " / ".join(f"{k}: {v}件" for k, v in dtype_log.items()),
        f"正常抽出: {len(records) - needs_rev_count} 件",
        f"要確認: {needs_rev_count} 件",
    ]
    for k, v in sorted(review_breakdown.items(), key=lambda x: -x[1]):
        log_lines.append(f"  ・{k}: {v} 件")
    if skipped_dup:
        log_lines.append(f"重複スキップ: {skipped_dup} 件")

    with open(os.path.join(outdir, "00_処理ログ.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join(log_lines))

    # マニフェストを更新（次回の差分処理のために全レコードを保存）
    # ※ needs_review=True のファイルはキャッシュに乗せない
    #   → 次回OCRありで再処理したとき、⚠ファイルだけが自動的に再処理される
    manifest_new: Dict[str, dict] = {"_cache_version": _CACHE_VERSION}
    for r in records:
        if r.sha1 and not r.needs_review:
            manifest_new[r.sha1] = asdict(r)
    try:
        with open(manifest_path, "w", encoding="utf-8") as f:
            json.dump(manifest_new, f, ensure_ascii=False, separators=(",", ":"))
    except Exception:
        pass  # マニフェスト保存失敗は致命的ではない

    breakdown_str = "　".join(f"{k}: {v}件" for k, v in sorted(review_breakdown.items(), key=lambda x: -x[1]))
    return len(records), needs_rev_count, breakdown_str
